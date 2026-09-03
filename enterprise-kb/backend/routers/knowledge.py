from __future__ import annotations

import asyncio
import io
import json
import logging
import os
import re
import time
import uuid
import zipfile
from typing import List, Optional
from urllib.parse import quote

import httpx
from fastapi import APIRouter, Depends, HTTPException, Query, Request, status
from fastapi.responses import Response, StreamingResponse
from open_webui.config import (
    BYPASS_ADMIN_ACCESS_CONTROL,
    OPENAI_API_BASE_URL,
    OPENAI_API_KEY,
    RAG_OPENAI_API_BASE_URL,
    RAG_OPENAI_API_KEY,
    RAG_TEMPLATE,
)
from open_webui.constants import ERROR_MESSAGES
from open_webui.events import EVENTS, publish_event
from open_webui.internal.db import get_async_session
from open_webui.models.access_grants import AccessGrants
from open_webui.models.config import Config
from open_webui.models.files import FileMetadataResponse, FileModel, FileModelResponse, Files
from open_webui.models.groups import Groups
from open_webui.models.knowledge import (
    KnowledgeChunkMergeForm,
    KnowledgeChunkModel,
    KnowledgeChunkPreviewForm,
    KnowledgeChunkSplitForm,
    KnowledgeChunk,
    KnowledgeDirectoryForm,
    KnowledgeDirectoryModel,
    KnowledgeEvaluateQueryForm,
    KnowledgeFile,
    KnowledgeFileListResponse,
    KnowledgeForm,
    KnowledgeProcessingTask,
    KnowledgeProcessingTaskModel,
    KnowledgeBatchTask,
    KnowledgeBatchTaskModel,
    KnowledgeRelevanceAnnotationForm,
    KnowledgeResponse,
    KnowledgeRelevanceJudgment,
    KnowledgeRelevanceJudgmentModel,
    KnowledgeSnapshot,
    KnowledgeSnapshotCompareForm,
    KnowledgeSnapshotCompareResult,
    KnowledgeSnapshotCreateForm,
    KnowledgeSnapshotModel,
    Knowledges,
    KnowledgeUserResponse,
)
from open_webui.models.models import ModelForm, Models
from open_webui.retrieval.vector.async_client import ASYNC_VECTOR_DB_CLIENT
from open_webui.retrieval.external import retrieve_external_knowledge, retrieve_external_knowledge_for_connection
from open_webui.retrieval.utils import (
    build_loader_from_config,
    get_embedding_function,
    get_loader_config,
    query_collection_with_hybrid_search,
    query_collection,
)
from open_webui.routers.retrieval import (
    BatchProcessFilesForm,
    ProcessFileForm,
    get_retrieval_config,
    process_file,
    process_files_batch,
    save_docs_to_vector_db,
)
from open_webui.storage.provider import Storage
from open_webui.utils.access_control import filter_allowed_access_grants, has_permission
from open_webui.utils.access_control.files import has_access_to_file
from open_webui.utils.auth import get_admin_user, get_verified_user
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession

log = logging.getLogger(__name__)

router = APIRouter()

# In-memory storage for workflow execution results (for Word download)
_exec_results = {}

############################
# getKnowledgeBases
############################

PAGE_ITEM_COUNT = 30

############################
# Knowledge Base Embedding
############################

# Knowledge that sits unread serves no one. Let what is
# stored here find the ones who need it.
KNOWLEDGE_BASES_COLLECTION = 'knowledge-bases'


async def embed_knowledge_base_metadata(
    request: Request,
    knowledge_base_id: str,
    name: str,
    description: str,
) -> bool:
    """Generate and store embedding for knowledge base."""
    try:
        content = f'{name}\n\n{description}' if description else name
        embedding = await request.app.state.EMBEDDING_FUNCTION(content)
        await ASYNC_VECTOR_DB_CLIENT.upsert(
            collection_name=KNOWLEDGE_BASES_COLLECTION,
            items=[
                {
                    'id': knowledge_base_id,
                    'text': content,
                    'vector': embedding,
                    'metadata': {
                        'knowledge_base_id': knowledge_base_id,
                    },
                }
            ],
        )
        return True
    except Exception as e:
        log.error(f'Failed to embed knowledge base {knowledge_base_id}: {e}')
        return False


async def remove_knowledge_base_metadata_embedding(knowledge_base_id: str) -> bool:
    """Remove knowledge base embedding."""
    try:
        await ASYNC_VECTOR_DB_CLIENT.delete(
            collection_name=KNOWLEDGE_BASES_COLLECTION,
            ids=[knowledge_base_id],
        )
        return True
    except Exception as e:
        log.debug(f'Failed to remove embedding for {knowledge_base_id}: {e}')
        return False


class KnowledgeAccessResponse(KnowledgeUserResponse):
    write_access: bool | None = False


class KnowledgeAccessListResponse(BaseModel):
    items: list[KnowledgeAccessResponse]
    total: int


def is_external_knowledge(knowledge) -> bool:
    return (knowledge.meta or {}).get('source') == 'external'


def external_knowledge_error():
    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail='External knowledge bases are read-only.',
    )


@router.get('/', response_model=KnowledgeAccessListResponse)
async def get_knowledge_bases(
    page: int | None = 1,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    page = max(page, 1)
    limit = PAGE_ITEM_COUNT
    skip = (page - 1) * limit

    filter = {}
    groups = await Groups.get_groups_by_member_id(user.id, db=db)
    user_group_ids = {group.id for group in groups}

    if not user.role == 'admin' or not BYPASS_ADMIN_ACCESS_CONTROL:
        if groups:
            filter['group_ids'] = [group.id for group in groups]

        filter['user_id'] = user.id

    result = await Knowledges.search_knowledge_bases(user.id, filter=filter, skip=skip, limit=limit, db=db)

    # Batch-fetch writable knowledge IDs in a single query instead of N has_access calls
    knowledge_base_ids = [knowledge_base.id for knowledge_base in result.items]
    writable_knowledge_base_ids = await AccessGrants.get_accessible_resource_ids(
        user_id=user.id,
        resource_type='knowledge',
        resource_ids=knowledge_base_ids,
        permission='write',
        user_group_ids=user_group_ids,
        db=db,
    )

    return KnowledgeAccessListResponse(
        items=[
            KnowledgeAccessResponse(
                **knowledge_base.model_dump(),
                write_access=(
                    user.id == knowledge_base.user_id
                    or (user.role == 'admin' and BYPASS_ADMIN_ACCESS_CONTROL)
                    or knowledge_base.id in writable_knowledge_base_ids
                ),
            )
            for knowledge_base in result.items
        ],
        total=result.total,
    )


@router.get('/search', response_model=KnowledgeAccessListResponse)
async def search_knowledge_bases(
    query: str | None = None,
    view_option: str | None = None,
    source: str | None = None,
    page: int | None = 1,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    page = max(page, 1)
    limit = PAGE_ITEM_COUNT
    skip = (page - 1) * limit

    filter = {}
    if query:
        filter['query'] = query
    if view_option:
        filter['view_option'] = view_option
    if source in {'local', 'external'}:
        filter['source'] = source

    groups = await Groups.get_groups_by_member_id(user.id, db=db)
    user_group_ids = {group.id for group in groups}

    if not user.role == 'admin' or not BYPASS_ADMIN_ACCESS_CONTROL:
        if groups:
            filter['group_ids'] = [group.id for group in groups]

        filter['user_id'] = user.id

    result = await Knowledges.search_knowledge_bases(user.id, filter=filter, skip=skip, limit=limit, db=db)

    # Batch-fetch writable knowledge IDs in a single query instead of N has_access calls
    knowledge_base_ids = [knowledge_base.id for knowledge_base in result.items]
    writable_knowledge_base_ids = await AccessGrants.get_accessible_resource_ids(
        user_id=user.id,
        resource_type='knowledge',
        resource_ids=knowledge_base_ids,
        permission='write',
        user_group_ids=user_group_ids,
        db=db,
    )

    return KnowledgeAccessListResponse(
        items=[
            KnowledgeAccessResponse(
                **knowledge_base.model_dump(),
                write_access=(
                    user.id == knowledge_base.user_id
                    or (user.role == 'admin' and BYPASS_ADMIN_ACCESS_CONTROL)
                    or knowledge_base.id in writable_knowledge_base_ids
                ),
            )
            for knowledge_base in result.items
        ],
        total=result.total,
    )


@router.get('/search/files', response_model=KnowledgeFileListResponse)
async def search_knowledge_files(
    query: str | None = None,
    include_content: bool = Query(False, description='Include file content in search (expensive).'),
    page: int | None = 1,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    page = max(page, 1)
    limit = PAGE_ITEM_COUNT
    skip = (page - 1) * limit

    filter = {}
    if query:
        filter['query'] = query
    if include_content:
        filter['include_content'] = True

    groups = await Groups.get_groups_by_member_id(user.id, db=db)
    if groups:
        filter['group_ids'] = [group.id for group in groups]

    filter['user_id'] = user.id

    return await Knowledges.search_knowledge_files(filter=filter, skip=skip, limit=limit, db=db)


############################
# CreateNewKnowledge
############################


@router.post('/create', response_model=KnowledgeResponse | None)
async def create_new_knowledge(
    request: Request,
    form_data: KnowledgeForm,
    user=Depends(get_verified_user),
):
    # NOTE: We intentionally do NOT use Depends(get_async_session) here.
    # Database operations (has_permission, filter_allowed_access_grants, insert_new_knowledge) manage their own sessions.
    # This prevents holding a connection during embed_knowledge_base_metadata()
    # which makes external embedding API calls (1-5+ seconds).
    if user.role != 'admin' and not await has_permission(
        user.id, 'workspace.knowledge', await Config.get('user.permissions')
    ):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail=ERROR_MESSAGES.UNAUTHORIZED,
        )

    form_data.access_grants = await filter_allowed_access_grants(
        await Config.get('user.permissions'),
        user.id,
        user.role,
        form_data.access_grants,
        'sharing.public_knowledge',
    )

    knowledge = await Knowledges.insert_new_knowledge(user.id, form_data)

    if knowledge:
        # Embed knowledge base for semantic search
        await embed_knowledge_base_metadata(
            request,
            knowledge.id,
            knowledge.name,
            knowledge.description,
        )
        await publish_event(
            request,
            EVENTS.KNOWLEDGE_CREATED,
            actor=user,
            subject_id=knowledge.id,
            data={'name': knowledge.name},
        )
        return knowledge
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.FILE_EXISTS,
        )


############################
# ReindexKnowledgeFiles
############################


@router.post('/reindex', response_model=bool)
async def reindex_knowledge_files(
    request: Request,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    if user.role != 'admin':
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail=ERROR_MESSAGES.UNAUTHORIZED,
        )

    knowledge_bases = await Knowledges.get_knowledge_bases(db=db)
    knowledge_base_files = [
        (knowledge_base, await Knowledges.get_files_by_id(knowledge_base.id, db=db))
        for knowledge_base in knowledge_bases
    ]
    total_files = sum(len(files) for _, files in knowledge_base_files)
    processed_files = 0
    failed_files = []
    start_time = time.monotonic()

    log.info(f'Starting reindexing for {len(knowledge_bases)} knowledge bases ({total_files} files)')

    for kb_idx, (knowledge_base, files) in enumerate(knowledge_base_files, start=1):
        try:
            try:
                if await ASYNC_VECTOR_DB_CLIENT.has_collection(collection_name=knowledge_base.id):
                    await ASYNC_VECTOR_DB_CLIENT.delete_collection(collection_name=knowledge_base.id)
            except Exception as e:
                log.error(f'Error deleting collection {knowledge_base.id}: {str(e)}')
                continue  # Skip, don't raise

            for file in files:
                processed_files += 1
                eta = ''
                if processed_files > 1:
                    elapsed = time.monotonic() - start_time
                    remaining_files = total_files - processed_files + 1
                    eta = f', ETA: {round(elapsed / (processed_files - 1) * remaining_files)}s'

                log.info(
                    f'Reindexing knowledge base {kb_idx}/{len(knowledge_bases)} '
                    f'file {processed_files}/{total_files}{eta}: {file.filename}'
                )

                try:
                    await process_file(
                        request,
                        ProcessFileForm(file_id=file.id, collection_name=knowledge_base.id),
                        user=user,
                        db=db,
                    )
                except Exception as e:
                    log.error(f'Error processing file {file.filename} (ID: {file.id}): {str(e)}')
                    failed_files.append({'file_id': file.id, 'error': str(e)})
                    continue

        except Exception as e:
            log.error(f'Error processing knowledge base {knowledge_base.id}: {str(e)}')
            # Don't raise, just continue
            continue

    if failed_files:
        log.warning(f'Failed to process {len(failed_files)} files')
        for failed in failed_files:
            log.warning(f'File ID: {failed["file_id"]}, Error: {failed["error"]}')

    log.info(f'Reindexing completed in {round(time.monotonic() - start_time)}s.')
    await publish_event(
        request,
        EVENTS.KNOWLEDGE_REINDEXED,
        actor=user,
        subject_id='all',
        data={'count': len(knowledge_bases)},
    )
    return True


############################
# ReindexKnowledgeBases
############################


@router.post('/metadata/reindex', response_model=dict)
async def reindex_knowledge_base_metadata_embeddings(
    request: Request,
    user=Depends(get_admin_user),
):
    """Batch embed all existing knowledge bases. Admin only.

    NOTE: We intentionally do NOT use Depends(get_async_session) here.
    This endpoint loops through ALL knowledge bases and calls embed_knowledge_base_metadata()
    for each one, making N external embedding API calls. Holding a session during
    this entire operation would exhaust the connection pool.
    """
    knowledge_bases = await Knowledges.get_knowledge_bases()
    log.info(f'Reindexing embeddings for {len(knowledge_bases)} knowledge bases')

    success_count = 0
    for kb in knowledge_bases:
        if await embed_knowledge_base_metadata(request, kb.id, kb.name, kb.description):
            success_count += 1

    log.info(f'Embedding reindex complete: {success_count}/{len(knowledge_bases)}')
    return {'total': len(knowledge_bases), 'success': success_count}


############################
# External Knowledge Sources
############################


class ExternalKnowledgeSourceForm(BaseModel):
    type: str = 'collection'
    name: str
    config: Optional[dict] = None


class ExternalKnowledgeCreateForm(BaseModel):
    name: str
    description: str = ''
    connection_id: str
    source: ExternalKnowledgeSourceForm
    access_grants: Optional[list[dict]] = None


class ExternalKnowledgeSourceCreateForm(BaseModel):
    name: str
    description: str = ''
    connection: ExternalKnowledgeConnectionForm
    source: ExternalKnowledgeSourceForm
    access_grants: Optional[list[dict]] = None
    test_query: str
    test_count: int = 5


class ExternalKnowledgeSourceUpdateForm(ExternalKnowledgeSourceCreateForm):
    pass


class ExternalKnowledgeSourceTestForm(BaseModel):
    connection_id: Optional[str] = None
    connection: ExternalKnowledgeConnectionForm
    source: ExternalKnowledgeSourceForm
    query: str
    count: int = 5


class ExternalKnowledgeRetrieveTestForm(BaseModel):
    query: str
    source: Optional[ExternalKnowledgeSourceForm] = None
    count: int = 5


class ExternalKnowledgeConnectionForm(BaseModel):
    name: str
    provider: str
    endpoint: str
    auth_config: Optional[dict] = None
    config: Optional[dict] = None
    capabilities: Optional[dict] = None
    enabled: bool = True


class ExternalKnowledgeConnectionListResponse(BaseModel):
    items: list[dict]
    total: int


EXTERNAL_KNOWLEDGE_CONNECTIONS_CONFIG_KEY = 'external_knowledge.connections'
EXTERNAL_KNOWLEDGE_PROVIDERS = {'qdrant', 'milvus', 'pgvector'}


def _validate_external_connection_form(form_data: ExternalKnowledgeConnectionForm) -> tuple[str, dict]:
    provider = form_data.provider.lower().strip()
    if provider not in EXTERNAL_KNOWLEDGE_PROVIDERS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail='Unsupported external knowledge provider.',
        )

    if not form_data.name.strip():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Knowledge source name is required.')

    if not form_data.endpoint.strip():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Knowledge source endpoint is required.')

    config = form_data.config or {}
    allowed_config_keys = {'timeout'}
    if provider == 'milvus':
        allowed_config_keys.add('db_name')

    return provider, {key: value for key, value in config.items() if key in allowed_config_keys}


def _external_auth_config(provider: str, incoming: Optional[dict], existing: Optional[dict] = None) -> dict:
    if provider == 'pgvector':
        return {}
    return existing if incoming is None else incoming or {}


def _normalize_external_source(source: ExternalKnowledgeSourceForm, provider: str) -> ExternalKnowledgeSourceForm:
    source.type = (source.type or 'collection').strip()
    source.name = source.name.strip()

    if source.type != 'collection':
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Only collection sources are supported.')
    if not source.name:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Collection name is required.')

    config = source.config or {}
    allowed_keys = {'content_field', 'metadata_field', 'document_id_field'}
    if provider in {'qdrant', 'milvus'}:
        allowed_keys.add('vector_field')
    if provider == 'pgvector':
        allowed_keys.update({'table_name', 'collection_field', 'vector_field'})

    normalized_config = {
        key: value.strip() if isinstance(value, str) else value
        for key, value in config.items()
        if key in allowed_keys and value is not None and (not isinstance(value, str) or value.strip())
    }

    if not normalized_config.get('content_field'):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Content field is required.')
    if provider in {'milvus', 'pgvector'} and not normalized_config.get('vector_field'):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Vector field is required.')

    source.config = normalized_config
    return source


def _sanitize_external_connection(connection: dict) -> dict:
    sanitized = {**connection}
    sanitized.pop('auth_config', None)
    sanitized['auth_configured'] = bool(connection.get('auth_config'))
    return sanitized


async def _get_external_connections() -> list[dict]:
    return await Config.get(EXTERNAL_KNOWLEDGE_CONNECTIONS_CONFIG_KEY, []) or []


async def _set_external_connections(connections: list[dict]) -> None:
    await Config.upsert({EXTERNAL_KNOWLEDGE_CONNECTIONS_CONFIG_KEY: connections})


def _external_connection_dict(
    form_data: ExternalKnowledgeConnectionForm, user_id: str, id: Optional[str] = None
) -> dict:
    provider, config = _validate_external_connection_form(form_data)
    now = int(time.time())
    return {
        'id': id or str(uuid.uuid4()),
        'name': form_data.name.strip(),
        'provider': provider,
        'endpoint': form_data.endpoint.strip(),
        'auth_config': _external_auth_config(provider, form_data.auth_config),
        'config': config,
        'capabilities': form_data.capabilities or {'retrieve': True},
        'health': None,
        'enabled': form_data.enabled,
        'created_by': user_id,
        'created_at': now,
        'updated_at': now,
    }


def _external_connection_update_dict(
    form_data: ExternalKnowledgeConnectionForm,
    existing: dict,
) -> dict:
    provider, config = _validate_external_connection_form(form_data)
    return {
        **existing,
        'name': form_data.name.strip(),
        'provider': provider,
        'endpoint': form_data.endpoint.strip(),
        'auth_config': _external_auth_config(provider, form_data.auth_config, existing.get('auth_config')) or {},
        'config': config,
        'capabilities': form_data.capabilities or {'retrieve': True},
        'enabled': form_data.enabled,
        'updated_at': int(time.time()),
    }


async def _get_external_connection(id: str) -> Optional[dict]:
    connections = await _get_external_connections()
    return next((connection for connection in connections if connection.get('id') == id), None)


async def _count_external_connection_mappings(connection_id: str, db: Optional[AsyncSession] = None) -> int:
    count = 0
    for knowledge in await Knowledges.get_knowledge_bases(db=db):
        if (knowledge.meta or {}).get('external', {}).get('connection_id') == connection_id:
            count += 1
    return count


@router.get('/external/connections', response_model=ExternalKnowledgeConnectionListResponse)
async def get_external_knowledge_connections(
    user=Depends(get_admin_user),
    db: AsyncSession = Depends(get_async_session),
):
    connections = [_sanitize_external_connection(connection) for connection in await _get_external_connections()]
    return ExternalKnowledgeConnectionListResponse(items=connections, total=len(connections))


@router.post('/external/connections', response_model=dict)
async def create_external_knowledge_connection(
    request: Request,
    form_data: ExternalKnowledgeConnectionForm,
    user=Depends(get_admin_user),
):
    connections = await _get_external_connections()
    connection = _external_connection_dict(form_data, user.id)
    connections.append(connection)
    await _set_external_connections(connections)
    sanitized = _sanitize_external_connection(connection)
    await publish_event(
        request,
        EVENTS.KNOWLEDGE_EXTERNAL_CONNECTION_CREATED,
        actor=user,
        subject_id=connection.get('id'),
        data={'name': sanitized.get('name'), 'provider': sanitized.get('provider')},
    )
    return sanitized


@router.get('/external/connections/{id}', response_model=dict)
async def get_external_knowledge_connection(
    id: str,
    user=Depends(get_admin_user),
    db: AsyncSession = Depends(get_async_session),
):
    connection = await _get_external_connection(id)
    if not connection:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)
    return _sanitize_external_connection(connection)


@router.patch('/external/connections/{id}', response_model=dict)
async def update_external_knowledge_connection(
    request: Request,
    id: str,
    form_data: ExternalKnowledgeConnectionForm,
    user=Depends(get_admin_user),
):
    connections = await _get_external_connections()
    idx = next((idx for idx, connection in enumerate(connections) if connection.get('id') == id), None)
    if idx is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    connection = _external_connection_update_dict(form_data, connections[idx])
    connections[idx] = connection
    await _set_external_connections(connections)
    sanitized = _sanitize_external_connection(connection)
    await publish_event(
        request,
        EVENTS.KNOWLEDGE_EXTERNAL_CONNECTION_UPDATED,
        actor=user,
        subject_id=id,
        data={'name': sanitized.get('name'), 'provider': sanitized.get('provider')},
    )
    return sanitized


@router.delete('/external/connections/{id}', response_model=bool)
async def delete_external_knowledge_connection(
    request: Request,
    id: str,
    user=Depends(get_admin_user),
    db: AsyncSession = Depends(get_async_session),
):
    connection = await _get_external_connection(id)
    if not connection:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    if await _count_external_connection_mappings(id, db=db) > 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail='External connection is still used by knowledge bases.',
        )

    connections = [connection for connection in await _get_external_connections() if connection.get('id') != id]
    await _set_external_connections(connections)
    await publish_event(
        request,
        EVENTS.KNOWLEDGE_EXTERNAL_CONNECTION_DELETED,
        actor=user,
        subject_id=id,
        data={'name': connection.get('name'), 'provider': connection.get('provider')},
    )
    return True


@router.post('/external/connections/{id}/test', response_model=dict)
async def test_external_knowledge_connection(
    id: str,
    user=Depends(get_admin_user),
    db: AsyncSession = Depends(get_async_session),
):
    connection = await _get_external_connection(id)
    if not connection:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    health = {
        'ok': bool(connection.get('enabled') and connection.get('endpoint')),
        'provider': connection.get('provider'),
        'checked_at': int(time.time()),
    }
    connections = await _get_external_connections()
    for item in connections:
        if item.get('id') == id:
            item['health'] = health
            item['updated_at'] = int(time.time())
            break
    await _set_external_connections(connections)
    return health


async def _test_external_source_definition(
    request: Request,
    connection: dict,
    source: ExternalKnowledgeSourceForm,
    query: str,
    count: int,
    user,
) -> dict:
    if not query.strip():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Test query is required.')

    source = _normalize_external_source(source, connection.get('provider'))
    test_knowledge = KnowledgeResponse(
        id='external-test',
        user_id=user.id,
        name=connection.get('name'),
        description='',
        meta={
            'source': 'external',
            'read_only': True,
            'external': {
                'connection_id': connection.get('id'),
                'source': source.model_dump(),
                'provider': connection.get('provider'),
                'auth_mode': 'service_account',
                'capabilities': {'retrieve': True},
            },
        },
        access_grants=[],
        created_at=int(time.time()),
        updated_at=int(time.time()),
    )
    result = await retrieve_external_knowledge_for_connection(
        request,
        test_knowledge,
        connection,
        [query.strip()],
        count,
        user=user,
    )
    return {
        'documents': result.get('documents', [[]])[0],
        'metadatas': result.get('metadatas', [[]])[0],
        'distances': result.get('distances', [[]])[0],
    }


@router.post('/external/source/test', response_model=dict)
async def test_external_knowledge_source(
    request: Request,
    form_data: ExternalKnowledgeSourceTestForm,
    user=Depends(get_admin_user),
):
    if form_data.connection_id:
        existing_connection = await _get_external_connection(form_data.connection_id)
        if not existing_connection:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail='External connection not found.')
        connection = _external_connection_update_dict(form_data.connection, existing_connection)
    else:
        connection = _external_connection_dict(form_data.connection, user.id, id='external-test')

    return await _test_external_source_definition(
        request,
        connection,
        form_data.source,
        form_data.query,
        form_data.count,
        user,
    )


@router.post('/external/connections/{id}/retrieve-test', response_model=dict)
async def test_external_knowledge_retrieval(
    request: Request,
    id: str,
    form_data: ExternalKnowledgeRetrieveTestForm,
    user=Depends(get_admin_user),
    db: AsyncSession = Depends(get_async_session),
):
    connection = await _get_external_connection(id)
    if not connection:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    source = form_data.source or ExternalKnowledgeSourceForm(name='test', config={'content_field': 'payload.text'})
    return await _test_external_source_definition(request, connection, source, form_data.query, form_data.count, user)


@router.post('/external/knowledge/create', response_model=KnowledgeResponse | None)
async def create_external_knowledge(
    request: Request,
    form_data: ExternalKnowledgeCreateForm,
    user=Depends(get_admin_user),
    db: AsyncSession = Depends(get_async_session),
):
    connection = await _get_external_connection(form_data.connection_id)
    if not connection:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)
    if not form_data.name.strip():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Knowledge name is required.')
    source = _normalize_external_source(form_data.source, connection.get('provider'))

    form_data.access_grants = await filter_allowed_access_grants(
        await Config.get('user.permissions'),
        user.id,
        user.role,
        form_data.access_grants,
        'sharing.public_knowledge',
    )

    knowledge = await Knowledges.insert_new_knowledge(
        user.id,
        KnowledgeForm(
            name=form_data.name.strip(),
            description=form_data.description,
            access_grants=form_data.access_grants,
        ),
        db=db,
    )
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=ERROR_MESSAGES.FILE_EXISTS)

    meta = {
        'source': 'external',
        'read_only': True,
        'external': {
            'connection_id': form_data.connection_id,
            'source': source.model_dump(),
            'provider': connection.get('provider'),
            'auth_mode': 'service_account',
            'capabilities': {'retrieve': True},
        },
    }
    knowledge = await Knowledges.update_knowledge_meta_by_id(knowledge.id, meta, db=db)
    await embed_knowledge_base_metadata(request, knowledge.id, knowledge.name, knowledge.description)
    return knowledge


@router.post('/external/source/create', response_model=KnowledgeResponse | None)
async def create_external_knowledge_source(
    request: Request,
    form_data: ExternalKnowledgeSourceCreateForm,
    user=Depends(get_admin_user),
    db: AsyncSession = Depends(get_async_session),
):
    if not form_data.name.strip():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Knowledge name is required.')

    connection = _external_connection_dict(form_data.connection, user.id)
    source = _normalize_external_source(form_data.source, connection.get('provider'))
    test_result = await _test_external_source_definition(
        request,
        connection,
        source,
        form_data.test_query,
        form_data.test_count,
        user,
    )
    if not test_result.get('documents'):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Test query returned no results.')

    form_data.access_grants = await filter_allowed_access_grants(
        await Config.get('user.permissions'),
        user.id,
        user.role,
        form_data.access_grants,
        'sharing.public_knowledge',
    )

    connections = await _get_external_connections()
    connections.append(connection)
    await _set_external_connections(connections)

    knowledge = await Knowledges.insert_new_knowledge(
        user.id,
        KnowledgeForm(
            name=form_data.name.strip(),
            description=form_data.description,
            access_grants=form_data.access_grants,
        ),
        db=db,
    )
    if not knowledge:
        connections = [item for item in await _get_external_connections() if item.get('id') != connection.get('id')]
        await _set_external_connections(connections)
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=ERROR_MESSAGES.FILE_EXISTS)

    meta = {
        'source': 'external',
        'read_only': True,
        'external': {
            'connection_id': connection.get('id'),
            'source': source.model_dump(),
            'provider': connection.get('provider'),
            'auth_mode': 'service_account',
            'capabilities': {'retrieve': True},
        },
    }
    knowledge = await Knowledges.update_knowledge_meta_by_id(knowledge.id, meta, db=db)
    await embed_knowledge_base_metadata(request, knowledge.id, knowledge.name, knowledge.description)
    return knowledge


@router.patch('/external/source/{id}', response_model=KnowledgeResponse | None)
async def update_external_knowledge_source(
    request: Request,
    id: str,
    form_data: ExternalKnowledgeSourceUpdateForm,
    user=Depends(get_admin_user),
    db: AsyncSession = Depends(get_async_session),
):
    knowledge = await Knowledges.get_knowledge_by_id(id=id, db=db)
    if not knowledge or not is_external_knowledge(knowledge):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)
    if not form_data.name.strip():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Knowledge name is required.')

    connection_id = (knowledge.meta or {}).get('external', {}).get('connection_id')
    connections = await _get_external_connections()
    idx = next((idx for idx, connection in enumerate(connections) if connection.get('id') == connection_id), None)
    if idx is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail='External connection not found.')

    existing_connection = connections[idx]
    connection = _external_connection_update_dict(form_data.connection, existing_connection)
    source = _normalize_external_source(form_data.source, connection.get('provider'))
    test_result = await _test_external_source_definition(
        request,
        connection,
        source,
        form_data.test_query,
        form_data.test_count,
        user,
    )
    if not test_result.get('documents'):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Test query returned no results.')

    form_data.access_grants = await filter_allowed_access_grants(
        await Config.get('user.permissions'),
        user.id,
        user.role,
        form_data.access_grants,
        'sharing.public_knowledge',
    )

    connections[idx] = connection
    await _set_external_connections(connections)

    updated = await Knowledges.update_knowledge_by_id(
        id=id,
        form_data=KnowledgeForm(
            name=form_data.name.strip(),
            description=form_data.description,
            access_grants=form_data.access_grants,
        ),
        db=db,
    )
    if not updated:
        connections[idx] = existing_connection
        await _set_external_connections(connections)
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=ERROR_MESSAGES.DEFAULT())

    meta = {
        'source': 'external',
        'read_only': True,
        'external': {
            'connection_id': connection.get('id'),
            'source': source.model_dump(),
            'provider': connection.get('provider'),
            'auth_mode': 'service_account',
            'capabilities': {'retrieve': True},
        },
    }
    updated = await Knowledges.update_knowledge_meta_by_id(id, meta, db=db)
    if not updated:
        connections[idx] = existing_connection
        await _set_external_connections(connections)
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=ERROR_MESSAGES.DEFAULT())

    await embed_knowledge_base_metadata(request, id, updated.name, updated.description)
    return updated


############################
# GetKnowledgeById
############################


class KnowledgeFilesResponse(KnowledgeResponse):
    files: list[FileMetadataResponse | None] = None
    write_access: bool | None = False


# -- Phase 8: Agent Workflow (appended) --

@router.get("/_workflows/roles")
async def get_agent_roles(user=Depends(get_verified_user)):
    from open_webui.models.knowledge import AGENT_ROLES_LIST
    return AGENT_ROLES_LIST

@router.get("/_workflows")
async def list_workflows(user=Depends(get_verified_user), db: AsyncSession = Depends(get_async_session)):
    from open_webui.models.knowledge import AgentWorkflow, AgentWorkflowStep, AgentWorkflowModel
    from sqlalchemy import select as sa_select
    r = await db.execute(sa_select(AgentWorkflow).where(AgentWorkflow.user_id == user.id).order_by(AgentWorkflow.updated_at.desc()))
    out = []
    for wf in r.scalars().all():
        sr = await db.execute(sa_select(AgentWorkflowStep).where(AgentWorkflowStep.workflow_id == wf.id).order_by(AgentWorkflowStep.order_index))
        steps = [{"id": s.id, "order_index": s.order_index, "agent_role": s.agent_role, "knowledge_id": s.knowledge_id, "prompt_template": s.prompt_template, "input_var": s.input_var, "output_var": s.output_var} for s in sr.scalars().all()]
        out.append(AgentWorkflowModel(id=wf.id, user_id=wf.user_id, name=wf.name, description=wf.description, steps=steps, created_at=wf.created_at, updated_at=wf.updated_at))
    return out

@router.post("/_workflows")
async def create_workflow(request: Request, user=Depends(get_verified_user), db: AsyncSession = Depends(get_async_session)):
    from open_webui.models.knowledge import AgentWorkflow, AgentWorkflowStep, AgentWorkflowModel
    b = await request.json()
    n = b.get("name", ""); d = b.get("description", ""); ss = b.get("steps", [])
    now = int(time.time())
    wf = AgentWorkflow(id=str(uuid.uuid4()), user_id=user.id, name=n, description=d, created_at=now, updated_at=now)
    db.add(wf)
    for s in ss:
        db.add(AgentWorkflowStep(id=str(uuid.uuid4()), workflow_id=wf.id, order_index=s.get("order_index", 0), agent_role=s["agent_role"], knowledge_id=s.get("knowledge_id"), prompt_template=s.get("prompt_template"), input_var=s.get("input_var"), output_var=s.get("output_var"), created_at=now))
    await db.commit()
    return AgentWorkflowModel(id=wf.id, user_id=wf.user_id, name=n, description=d, steps=ss, created_at=now, updated_at=now)

@router.post("/_workflows/exec")
@router.post("/_workflows/exec")
async def exec_workflow(request: Request, user=Depends(get_verified_user), db: AsyncSession = Depends(get_async_session)):
    from open_webui.models.knowledge import AgentWorkflow, AgentWorkflowStep, AGENT_ROLES
    from sqlalchemy import select as sa_select
    b = await request.json()
    wid = b.get("workflow_id", ""); q = b.get("query", "")
    wf = (await db.execute(sa_select(AgentWorkflow).where(AgentWorkflow.id == wid))).scalars().first()
    if not wf: raise HTTPException(status_code=404, detail="Workflow not found")
    sr = await db.execute(sa_select(AgentWorkflowStep).where(AgentWorkflowStep.workflow_id == wf.id).order_by(AgentWorkflowStep.order_index))
    steps = sr.scalars().all()
    run_id = str(uuid.uuid4())
    results = []; variables = {"query": q}
    SSE = chr(10) + chr(10)
    SEP = chr(10) + chr(10) + "---" + chr(10) + chr(10)

    from dotenv import load_dotenv; load_dotenv(override=True); import os; OPENAI_API_BASE_URL = os.getenv("OPENAI_API_BASE_URL", "https://api.deepseek.com/v1"); OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
    api_url = OPENAI_API_BASE_URL.rstrip("/") if OPENAI_API_BASE_URL else "https://api.deepseek.com/v1"
    api_key = OPENAI_API_KEY or ""

    async def event_generator():
        for step in steps:
            role_info = AGENT_ROLES.get(step.agent_role, {"name": step.agent_role})
            yield "data: " + json.dumps({"step": step.order_index, "role": role_info["name"], "status": "running"}) + SSE
            output = ""
            if step.agent_role == "retriever" and step.knowledge_id:
                try:
                    emb_fn = request.app.state.EMBEDDING_FUNCTION
                    from open_webui.retrieval.utils import query_collection
                    r = await query_collection(request=request, collection_names=[step.knowledge_id], queries=[q], embedding_function=emb_fn, k=15)
                    if isinstance(r, dict) and r.get("documents") and r["documents"][0]:
                        docs = r["documents"][0][:15]
                        metas = r.get("metadatas", [[]])[0][:15] if r.get("metadatas") else []
                        # Group chunks by file_id and concatenate
                        from collections import OrderedDict
                        file_groups = OrderedDict()
                        for i, doc in enumerate(docs):
                            if not doc: continue
                            fid = ""
                            ci = i
                            if i < len(metas) and metas[i]:
                                fid = metas[i].get("file_id", "")
                                ci = metas[i].get("chunk_index", i)
                            if fid not in file_groups:
                                file_groups[fid] = []
                            file_groups[fid].append((ci, doc))
                        output = "[Retrieved " + str(len(docs)) + " chunks from " + str(len(file_groups)) + " documents]" + chr(10) + chr(10)
                        doc_num = 0
                        for fid, chunks in file_groups.items():
                            doc_num += 1
                            # Sort by chunk_index
                            chunks.sort(key=lambda x: x[0])
                            # Get filename
                            src = fid[:20] if fid else "Unknown"
                            if fid:
                                try:
                                    from open_webui.models.files import Files
                                    f = await Files.get_file_by_id(fid)
                                    if f:
                                        src = f.filename
                                except Exception:
                                    pass
                            # Concatenate chunks
                            merged = chr(10) + chr(10).join(c[1] for c in chunks)
                            output += chr(10) + "=== Document " + str(doc_num) + ": " + str(src) + " ===" + chr(10)
                            output += merged[:8000] + chr(10)
                    else:
                        output = "(no documents found in knowledge base)"
                except Exception as e:
                    output = "(retrieval error: " + str(e)[:100] + ")"
            else:
                try:
                    prompt_text = step.prompt_template or role_info.get("default_prompt", "")
                    # Collect ALL previous step outputs as context
                    context_parts = []
                    for vn, vv in variables.items():
                        if vn != "query" and vv and isinstance(vv, str) and len(vv) > 5:
                            context_parts.append("[" + vn + "]: " + str(vv)[:6000])
                    prev_context = chr(10).join(context_parts) if context_parts else ""
                    # Build final prompt — ALWAYS include the user's question first
                    prompt_text = "用户问题：" + q + chr(10) + chr(10)
                    if prev_context:
                        prompt_text += "前面步骤收集到的信息：" + chr(10) + prev_context + chr(10) + chr(10)
                    prompt_text += "你的任务：" + role_info.get("default_prompt", "")
                    prompt_text = prompt_text.replace("{context}", prev_context)
                    prompt_text = prompt_text.replace("{query}", q)
                    for vn, vv in variables.items():
                        prompt_text = prompt_text.replace("{" + vn + "}", str(vv)[:8000])
                    if api_key:
                        import httpx
                        async with httpx.AsyncClient(timeout=60) as client:
                            resp = await client.post(
                                api_url + "/chat/completions",
                                headers={"Authorization": "Bearer " + api_key},
                                json={"model": "deepseek-chat", "messages": [
                                    {"role": "system", "content": "You are the " + role_info["name"] + ". Respond in Chinese."},
                                    {"role": "user", "content": prompt_text[:10000]}
                                ], "max_tokens": 2048}
                            )
                            if resp.status_code == 200:
                                output = resp.json()["choices"][0]["message"]["content"]
                            else:
                                output = "(LLM HTTP " + str(resp.status_code) + ")"
                    else:
                        output = "(No API key - would call " + role_info["name"] + ")"
                except Exception as e:
                    output = "(LLM error: " + str(e)[:150] + ")"
            if step.output_var:
                variables[step.output_var] = output
            results.append({"step": step.order_index, "role": role_info["name"], "status": "done", "output": output})
            # SSE preview: truncated for display
            yield "data: " + json.dumps({"step": step.order_index, "role": role_info["name"], "status": "done", "output": output[:600]}) + SSE
        # Store full results for download
        from open_webui.routers.knowledge import _exec_results
        _exec_results[run_id] = {"query": q, "wf_name": wf.name, "steps": results, "ts": time.time()}
        yield "data: " + json.dumps({"status": "done", "run_id": run_id, "results": results}) + SSE
    return StreamingResponse(event_generator(), media_type="text/event-stream")


@router.post("/_workflows/exec/download")
async def download_workflow_results(request: Request):
    """Generate and download a Word document from workflow execution results."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    b = await request.json()
    run_id = b.get("run_id", "")
    wf_name = b.get("wf_name", "Workflow")
    query = b.get("query", "")
    steps = b.get("steps", [])
    # Try to get full results from memory if run_id provided
    if run_id and run_id in _exec_results:
        cached = _exec_results[run_id]
        steps = cached.get("steps", steps)
        wf_name = cached.get("wf_name", wf_name)
        query = cached.get("query", query)
    doc = Document()
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    # Title
    title = doc.add_heading(f'{wf_name} - Execution Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Metadata
    doc.add_paragraph(f'Query: {query}')
    from datetime import datetime
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')
    # Steps
    for i, s in enumerate(steps):
        role = s.get("role", f"Step {i+1}")
        output = s.get("output", "")
        doc.add_heading(f'{i+1}. {role}', level=2)
        # Split output into paragraphs
        for line in output.split(chr(10)):
            line = line.strip()
            if line:
                doc.add_paragraph(line)
        doc.add_paragraph('')
    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    # Cleanup old entries
    now_ts = time.time()
    for k in list(_exec_results.keys()):
        if now_ts - _exec_results[k].get("ts", 0) > 3600:
            del _exec_results[k]
    filename = f'{wf_name}_report.docx'
    return Response(
        content=buf.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename*=UTF-8\'\'{quote(filename)}'}
    )


@router.delete("/_workflows/{wfid}")
async def del_workflow(wfid: str, user=Depends(get_verified_user), db: AsyncSession = Depends(get_async_session)):
    from open_webui.models.knowledge import AgentWorkflow, AgentWorkflowStep
    from sqlalchemy import delete as sa_delete
    await db.execute(sa_delete(AgentWorkflowStep).where(AgentWorkflowStep.workflow_id == wfid))
    await db.execute(sa_delete(AgentWorkflow).where(AgentWorkflow.id == wfid, AgentWorkflow.user_id == user.id))
    await db.commit()
    return {"status": True}
@router.get('/{id}', response_model=KnowledgeFilesResponse | None)
async def get_knowledge_by_id(id: str, user=Depends(get_verified_user), db: AsyncSession = Depends(get_async_session)):
    knowledge = await Knowledges.get_knowledge_by_id(id=id, db=db)

    if knowledge:
        if (
            user.role == 'admin'
            or knowledge.user_id == user.id
            or await AccessGrants.has_access(
                user_id=user.id,
                resource_type='knowledge',
                resource_id=knowledge.id,
                permission='read',
                db=db,
            )
        ):
            return KnowledgeFilesResponse(
                **knowledge.model_dump(),
                write_access=(
                    user.id == knowledge.user_id
                    or (user.role == 'admin' and BYPASS_ADMIN_ACCESS_CONTROL)
                    or await AccessGrants.has_access(
                        user_id=user.id,
                        resource_type='knowledge',
                        resource_id=knowledge.id,
                        permission='write',
                        db=db,
                    )
                ),
            )
        else:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail=ERROR_MESSAGES.ACCESS_PROHIBITED,
            )
    else:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )


############################
# UpdateKnowledgeById
############################


@router.post('/{id}/update', response_model=KnowledgeFilesResponse | None)
async def update_knowledge_by_id(
    request: Request,
    id: str,
    form_data: KnowledgeForm,
    user=Depends(get_verified_user),
):
    # NOTE: We intentionally do NOT use Depends(get_async_session) here.
    # Database operations manage their own short-lived sessions internally.
    # This prevents holding a connection during embed_knowledge_base_metadata()
    # which makes external embedding API calls (1-5+ seconds).
    knowledge = await Knowledges.get_knowledge_by_id(id=id)
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )
    # Is the user the original creator, in a group with write access, or an admin
    if (
        knowledge.user_id != user.id
        and not await AccessGrants.has_access(
            user_id=user.id,
            resource_type='knowledge',
            resource_id=knowledge.id,
            permission='write',
        )
        and user.role != 'admin'
    ):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.ACCESS_PROHIBITED,
        )

    form_data.access_grants = await filter_allowed_access_grants(
        await Config.get('user.permissions'),
        user.id,
        user.role,
        form_data.access_grants,
        'sharing.public_knowledge',
    )

    knowledge = await Knowledges.update_knowledge_by_id(id=id, form_data=form_data)
    if knowledge:
        # Re-embed knowledge base for semantic search
        await embed_knowledge_base_metadata(
            request,
            knowledge.id,
            knowledge.name,
            knowledge.description,
        )
        response = KnowledgeFilesResponse(
            **knowledge.model_dump(),
            files=await Knowledges.get_file_metadatas_by_id(knowledge.id),
        )
        await publish_event(
            request,
            EVENTS.KNOWLEDGE_UPDATED,
            actor=user,
            subject_id=knowledge.id,
            data={'name': knowledge.name},
        )
        return response
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.ID_TAKEN,
        )


############################
# UpdateKnowledgeAccessById
############################


class KnowledgeAccessGrantsForm(BaseModel):
    access_grants: list[dict]


@router.post('/{id}/access/update', response_model=KnowledgeFilesResponse | None)
async def update_knowledge_access_by_id(
    request: Request,
    id: str,
    form_data: KnowledgeAccessGrantsForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    knowledge = await Knowledges.get_knowledge_by_id(id=id, db=db)
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    if (
        knowledge.user_id != user.id
        and not await AccessGrants.has_access(
            user_id=user.id,
            resource_type='knowledge',
            resource_id=knowledge.id,
            permission='write',
            db=db,
        )
        and user.role != 'admin'
    ):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.ACCESS_PROHIBITED,
        )

    form_data.access_grants = await filter_allowed_access_grants(
        await Config.get('user.permissions'),
        user.id,
        user.role,
        form_data.access_grants,
        'sharing.public_knowledge',
    )

    knowledge.access_grants = await AccessGrants.set_access_grants('knowledge', id, form_data.access_grants, db=db)

    response = KnowledgeFilesResponse(
        **knowledge.model_dump(),
        files=await Knowledges.get_file_metadatas_by_id(id, db=db),
    )
    await publish_event(
        request,
        EVENTS.KNOWLEDGE_ACCESS_UPDATED,
        actor=user,
        subject_id=knowledge.id,
        data={'name': knowledge.name},
    )
    return response


############################
# GetPendingKnowledgeFiles
############################


@router.get('/{id}/files/pending')
async def get_pending_knowledge_files(
    id: str,
    stream: bool = Query(False),
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Return files that are being processed for this knowledge base but not yet linked.

    After a file is uploaded with ``knowledge_id`` in its metadata, the backend
    processes it in a background task before linking it to the ``knowledge_file``
    join table.  During this window the file is invisible to the normal file
    list endpoint.  This endpoint exposes those in-flight files so the frontend
    can show them with a processing indicator even after a page reload.

    When ``stream=true``, returns an SSE stream that polls every 3 seconds
    and emits the current pending file list.  Closes when no files remain.
    """
    knowledge = await Knowledges.get_knowledge_by_id(id=id, db=db)
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    if not (
        user.role == 'admin'
        or knowledge.user_id == user.id
        or await AccessGrants.has_access(
            user_id=user.id,
            resource_type='knowledge',
            resource_id=knowledge.id,
            permission='read',
            db=db,
        )
    ):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.ACCESS_PROHIBITED,
        )

    if not stream:
        return await Files.get_pending_files_for_knowledge(id, db=db)

    async def event_stream(knowledge_id: str):
        MAX_POLL_DURATION = 3600  # 1 hour max
        for _ in range(MAX_POLL_DURATION // 3):
            pending = await Files.get_pending_files_for_knowledge(knowledge_id)
            data = [f.model_dump() for f in pending]
            yield f'data: {json.dumps(data)}\n\n'
            if len(pending) == 0:
                break
            await asyncio.sleep(3)

    return StreamingResponse(
        event_stream(id),
        media_type='text/event-stream',
    )


############################
# GetKnowledgeFilesById
############################


@router.get('/{id}/files', response_model=KnowledgeFileListResponse)
async def get_knowledge_files_by_id(
    id: str,
    query: str | None = None,
    include_content: bool = Query(False, description='Include file content in search (expensive).'),
    view_option: str | None = None,
    order_by: str | None = None,
    direction: str | None = None,
    directory_id: str | None = Query(None, description='Filter by directory ID. Pass empty string for root.'),
    page: int | None = 1,
    limit: int | None = Query(None, description='Page size (admin only). Defaults to 30.'),
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    knowledge = await Knowledges.get_knowledge_by_id(id=id, db=db)
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    if not (
        user.role == 'admin'
        or knowledge.user_id == user.id
        or await AccessGrants.has_access(
            user_id=user.id,
            resource_type='knowledge',
            resource_id=knowledge.id,
            permission='read',
            db=db,
        )
    ):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.ACCESS_PROHIBITED,
        )

    page = max(page, 1)

    # Allow admins to configure page size; non-admins always get the default
    if user.role == 'admin' and limit is not None:
        limit = max(1, limit)
    else:
        limit = PAGE_ITEM_COUNT
    skip = (page - 1) * limit

    filter = {}
    if query:
        filter['query'] = query
    if include_content:
        filter['include_content'] = True
    if view_option:
        filter['view_option'] = view_option
    if order_by:
        filter['order_by'] = order_by
    if direction:
        filter['direction'] = direction
    # directory_id filtering: present in filter = scope to that directory (None = root)
    if directory_id is not None:
        filter['directory_id'] = directory_id if directory_id else None

    return await Knowledges.search_files_by_id(id, user.id, filter=filter, skip=skip, limit=limit, db=db)


############################
# AddFileToKnowledge
############################


class KnowledgeFileIdForm(BaseModel):
    file_id: str
    directory_id: Optional[str] = None


@router.post('/{id}/file/add', response_model=KnowledgeFilesResponse | None)
async def add_file_to_knowledge_by_id(
    request: Request,
    id: str,
    form_data: KnowledgeFileIdForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    knowledge = await Knowledges.get_knowledge_by_id(id=id, db=db)
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )
    if is_external_knowledge(knowledge):
        external_knowledge_error()

    if (
        knowledge.user_id != user.id
        and not await AccessGrants.has_access(
            user_id=user.id,
            resource_type='knowledge',
            resource_id=knowledge.id,
            permission='write',
            db=db,
        )
        and user.role != 'admin'
    ):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.ACCESS_PROHIBITED,
        )

    file = await Files.get_file_by_id(form_data.file_id, db=db)
    if not file:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )
    if not file.data:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.FILE_NOT_PROCESSED,
        )

    # KB write-access alone is not enough — caller must also be able to read the file.
    if file.user_id != user.id and user.role != 'admin':
        if not await has_access_to_file(file.id, 'read', user, db=db):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail=ERROR_MESSAGES.ACCESS_PROHIBITED,
            )

    # Add content to the vector database
    try:
        await _update_processing_progress(
            knowledge_id=id, file_id=form_data.file_id,
            status='embedding', progress_pct=50, db=db,
        )

        await process_file(
            request,
            ProcessFileForm(file_id=form_data.file_id, collection_name=id),
            user=user,
            db=db,
        )

        await _update_processing_progress(
            knowledge_id=id, file_id=form_data.file_id,
            status='completed', progress_pct=100, db=db,
        )

        # ── Automatically persist chunks to knowledge_chunk table ──
        try:
            from sqlalchemy import delete as sa_delete
            result_vdb = await ASYNC_VECTOR_DB_CLIENT.query(
                collection_name=id, filter={'file_id': form_data.file_id}
            )
            if result_vdb and result_vdb.ids and len(result_vdb.ids) > 0:
                import hashlib
                now = int(time.time())
                # Remove old chunk records
                await db.execute(
                    sa_delete(KnowledgeChunk).where(
                        KnowledgeChunk.knowledge_id == id,
                        KnowledgeChunk.file_id == form_data.file_id,
                    )
                )
                for idx, (doc_id, text, metadata) in enumerate(
                    zip(result_vdb.ids[0], result_vdb.documents[0], result_vdb.metadatas[0])
                ):
                    content_hash = hashlib.sha256(text.encode()).hexdigest() if text else None
                    chunk = KnowledgeChunk(
                        id=str(uuid.uuid4()),
                        knowledge_id=id,
                        file_id=form_data.file_id,
                        chunk_index=idx,
                        content=text,
                        token_count=len(text) // 4 if text else None,
                        meta=metadata,
                        content_hash=content_hash,
                        created_at=now,
                        updated_at=now,
                    )
                    db.add(chunk)
                await db.commit()
        except Exception as chunk_err:
            log.warning(f'Failed to persist chunks for KB {id} file {form_data.file_id}: {chunk_err}')

        # Add file to knowledge base
        await Knowledges.add_file_to_knowledge_by_id(
            knowledge_id=id,
            file_id=form_data.file_id,
            user_id=user.id,
            directory_id=form_data.directory_id,
            db=db,
        )
    except Exception as e:
        log.debug(e)
        await _update_processing_progress(
            knowledge_id=id, file_id=form_data.file_id,
            status='failed', error_message=str(e), db=db,
        )
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e),
        )

    if knowledge:
        response = KnowledgeFilesResponse(
            **knowledge.model_dump(),
            files=await Knowledges.get_file_metadatas_by_id(knowledge.id, db=db),
        )
        await publish_event(
            request,
            EVENTS.KNOWLEDGE_FILE_ADDED,
            actor=user,
            subject_id=form_data.file_id,
            data={'knowledge_id': knowledge.id, 'directory_id': form_data.directory_id},
        )
        return response
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )


@router.post('/{id}/file/update', response_model=KnowledgeFilesResponse | None)
async def update_file_from_knowledge_by_id(
    request: Request,
    id: str,
    form_data: KnowledgeFileIdForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    knowledge = await Knowledges.get_knowledge_by_id(id=id, db=db)
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )
    if is_external_knowledge(knowledge):
        external_knowledge_error()

    if (
        knowledge.user_id != user.id
        and not await AccessGrants.has_access(
            user_id=user.id,
            resource_type='knowledge',
            resource_id=knowledge.id,
            permission='write',
            db=db,
        )
        and user.role != 'admin'
    ):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.ACCESS_PROHIBITED,
        )

    file = await Files.get_file_by_id(form_data.file_id, db=db)
    if not file:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    # Validate the file actually belongs to this knowledge base
    if not await Knowledges.has_file(knowledge_id=id, file_id=form_data.file_id, db=db):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    # Remove content from the vector database
    await ASYNC_VECTOR_DB_CLIENT.delete(collection_name=knowledge.id, filter={'file_id': form_data.file_id})

    # Add content to the vector database
    try:
        await process_file(
            request,
            ProcessFileForm(file_id=form_data.file_id, collection_name=id),
            user=user,
            db=db,
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e),
        )

    if knowledge:
        response = KnowledgeFilesResponse(
            **knowledge.model_dump(),
            files=await Knowledges.get_file_metadatas_by_id(knowledge.id, db=db),
        )
        await publish_event(
            request,
            EVENTS.KNOWLEDGE_FILE_UPDATED,
            actor=user,
            subject_id=form_data.file_id,
            data={'knowledge_id': knowledge.id},
        )
        return response
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )


############################
# RemoveFileFromKnowledge
############################


@router.post('/{id}/file/remove', response_model=KnowledgeFilesResponse | None)
async def remove_file_from_knowledge_by_id(
    request: Request,
    id: str,
    form_data: KnowledgeFileIdForm,
    delete_file: bool = Query(True),
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    knowledge = await Knowledges.get_knowledge_by_id(id=id, db=db)
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )
    if is_external_knowledge(knowledge):
        external_knowledge_error()

    if (
        knowledge.user_id != user.id
        and not await AccessGrants.has_access(
            user_id=user.id,
            resource_type='knowledge',
            resource_id=knowledge.id,
            permission='write',
            db=db,
        )
        and user.role != 'admin'
    ):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.ACCESS_PROHIBITED,
        )

    file = await Files.get_file_by_id(form_data.file_id, db=db)
    if not file:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    # Validate the file actually belongs to this knowledge base
    if not await Knowledges.has_file(knowledge_id=id, file_id=form_data.file_id, db=db):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    await Knowledges.remove_file_from_knowledge_by_id(knowledge_id=id, file_id=form_data.file_id, db=db)

    # Remove content from the vector database
    try:
        await ASYNC_VECTOR_DB_CLIENT.delete(
            collection_name=knowledge.id, filter={'file_id': form_data.file_id}
        )  # Remove by file_id first

        await ASYNC_VECTOR_DB_CLIENT.delete(
            collection_name=knowledge.id, filter={'hash': file.hash}
        )  # Remove by hash as well in case of duplicates
    except Exception as e:
        log.debug('This was most likely caused by bypassing embedding processing')
        log.debug(e)
        pass

    # Anyone with write permission or higher can delete files
    if delete_file and (file.user_id == user.id or user.role == 'admin'):
        try:
            # Remove the file's collection from vector database
            file_collection = f'file-{form_data.file_id}'
            if await ASYNC_VECTOR_DB_CLIENT.has_collection(collection_name=file_collection):
                await ASYNC_VECTOR_DB_CLIENT.delete_collection(collection_name=file_collection)
        except Exception as e:
            log.debug('This was most likely caused by bypassing embedding processing')
            log.debug(e)
            pass

        # Delete file from database
        await Files.delete_file_by_id(form_data.file_id, db=db)

    if knowledge:
        response = KnowledgeFilesResponse(
            **knowledge.model_dump(),
            files=await Knowledges.get_file_metadatas_by_id(knowledge.id, db=db),
        )
        await publish_event(
            request,
            EVENTS.KNOWLEDGE_FILE_REMOVED,
            actor=user,
            subject_id=form_data.file_id,
            data={'knowledge_id': knowledge.id, 'delete_file': delete_file},
        )
        return response
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )


############################
# DeleteKnowledgeById
############################


@router.delete('/{id}/delete', response_model=bool)
async def delete_knowledge_by_id(
    request: Request,
    id: str,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    knowledge = await Knowledges.get_knowledge_by_id(id=id, db=db)
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    if (
        knowledge.user_id != user.id
        and not await AccessGrants.has_access(
            user_id=user.id,
            resource_type='knowledge',
            resource_id=knowledge.id,
            permission='write',
            db=db,
        )
        and user.role != 'admin'
    ):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.ACCESS_PROHIBITED,
        )

    log.info(f'Deleting knowledge base: {id} (name: {knowledge.name})')

    # Get all models
    models = await Models.get_all_models(db=db)
    log.info(f'Found {len(models)} models to check for knowledge base {id}')

    # Update models that reference this knowledge base
    for model in models:
        if model.meta and hasattr(model.meta, 'knowledge'):
            knowledge_list = model.meta.knowledge or []
            # Filter out the deleted knowledge base
            updated_knowledge = [k for k in knowledge_list if k.get('id') != id]

            # If the knowledge list changed, update the model
            if len(updated_knowledge) != len(knowledge_list):
                log.info(f'Updating model {model.id} to remove knowledge base {id}')
                model.meta.knowledge = updated_knowledge
                model_form = ModelForm(**model.model_dump())
                await Models.update_model_by_id(model.id, model_form, db=db)

    # Clean up vector DB
    if is_external_knowledge(knowledge):
        connection_id = (knowledge.meta or {}).get('external', {}).get('connection_id')
        if connection_id:
            connections = [
                connection for connection in await _get_external_connections() if connection.get('id') != connection_id
            ]
            await _set_external_connections(connections)
    else:
        try:
            await ASYNC_VECTOR_DB_CLIENT.delete_collection(collection_name=id)
        except Exception as e:
            log.debug(e)
            pass

    # Remove knowledge base embedding
    await remove_knowledge_base_metadata_embedding(id)

    result = await Knowledges.delete_knowledge_by_id(id=id, db=db)
    if result:
        await publish_event(
            request,
            EVENTS.KNOWLEDGE_DELETED,
            actor=user,
            subject_id=id,
            data={'name': knowledge.name},
        )
    return result


############################
# ResetKnowledgeById
############################


@router.post('/{id}/reset', response_model=KnowledgeResponse | None)
async def reset_knowledge_by_id(
    request: Request,
    id: str,
    include_directories: bool = Query(True),
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    knowledge = await Knowledges.get_knowledge_by_id(id=id, db=db)
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )
    if is_external_knowledge(knowledge):
        external_knowledge_error()

    if (
        knowledge.user_id != user.id
        and not await AccessGrants.has_access(
            user_id=user.id,
            resource_type='knowledge',
            resource_id=knowledge.id,
            permission='write',
            db=db,
        )
        and user.role != 'admin'
    ):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.ACCESS_PROHIBITED,
        )

    try:
        await ASYNC_VECTOR_DB_CLIENT.delete_collection(collection_name=id)
    except Exception as e:
        log.debug(e)
        pass

    knowledge = await Knowledges.reset_knowledge_by_id(id=id, include_directories=include_directories, db=db)
    if knowledge:
        await publish_event(
            request,
            EVENTS.KNOWLEDGE_RESET,
            actor=user,
            subject_id=id,
            data={'include_directories': include_directories},
        )
    return knowledge


############################
# SyncKnowledgeDiff
############################


class FileManifestEntry(BaseModel):
    filename: str  # basename: "readme.md"
    path: str  # relative dir: "docs/api" or "" for root
    checksum: str  # SHA-256 of raw bytes
    size: int


class SyncDiffForm(BaseModel):
    manifest: list[FileManifestEntry]


class SyncDiffResponse(BaseModel):
    added: list[dict]  # [{filename, path}] — new files
    modified: list[dict]  # [{filename, path, stale_file_id}] — changed files
    deleted: list[dict]  # [{file_id, filename}] — files to remove
    mkdir: list[str]  # directory paths to create
    rmdir: list[str]  # directory IDs to remove
    unmodified_count: int
    directory_map: dict[str, str]  # existing path → directory ID


@router.post('/{id}/sync/diff', response_model=SyncDiffResponse)
async def sync_knowledge_diff(
    id: str,
    form_data: SyncDiffForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """
    Compare a local file manifest against the knowledge base to determine
    which files need uploading, removing, and which directories to create/remove.
    """
    await _verify_knowledge_write_access(id, user, db)

    # ── Index existing state ──
    knowledge_files = await Knowledges.get_files_with_directory_ids(id, db=db)
    existing_directories = await Knowledges.get_all_directories(id, db=db)

    # Build directory path lookups
    directory_path_by_id: dict[str, str] = {}
    directory_id_by_path: dict[str, str] = {}
    for directory in existing_directories:
        segments = [directory.name]
        parent_id = directory.parent_id
        while parent_id:
            parent = next((d for d in existing_directories if d.id == parent_id), None)
            if not parent:
                break
            segments.insert(0, parent.name)
            parent_id = parent.parent_id
        full_path = '/'.join(segments)
        directory_path_by_id[directory.id] = full_path
        directory_id_by_path[full_path] = directory.id

    # Index existing files by (path, filename) → {file_id, checksum}
    indexed_files: dict[tuple[str, str], dict] = {}
    for file_model, directory_id in knowledge_files:
        file_path = directory_path_by_id.get(directory_id, '') if directory_id else ''
        stored_checksum = (file_model.meta or {}).get('file_hash')
        indexed_files[(file_path, file_model.filename)] = {
            'file_id': file_model.id,
            'checksum': stored_checksum,
        }

    # ── Diff files ──
    added: list[dict] = []
    modified: list[dict] = []
    deleted: list[dict] = []
    unmodified_count = 0
    manifest_keys: set[tuple[str, str]] = set()

    for entry in form_data.manifest:
        key = (entry.path, entry.filename)
        manifest_keys.add(key)

        if key not in indexed_files:
            added.append({'filename': entry.filename, 'path': entry.path})
        elif indexed_files[key]['checksum'] != entry.checksum:
            modified.append(
                {
                    'filename': entry.filename,
                    'path': entry.path,
                    'stale_file_id': indexed_files[key]['file_id'],
                }
            )
        else:
            unmodified_count += 1

    for key, file_info in indexed_files.items():
        if key not in manifest_keys:
            deleted.append({'file_id': file_info['file_id'], 'filename': key[1]})

    # ── Diff directories ──
    required_directory_paths: set[str] = set()
    for entry in form_data.manifest:
        if entry.path:
            segments = entry.path.split('/')
            for depth in range(len(segments)):
                required_directory_paths.add('/'.join(segments[: depth + 1]))

    mkdir = sorted([p for p in required_directory_paths if p not in directory_id_by_path], key=lambda p: p.count('/'))

    orphaned_directory_paths = set(directory_id_by_path) - required_directory_paths
    rmdir = [directory_id_by_path[p] for p in orphaned_directory_paths]

    return SyncDiffResponse(
        added=added,
        modified=modified,
        deleted=deleted,
        mkdir=mkdir,
        rmdir=rmdir,
        unmodified_count=unmodified_count,
        directory_map=directory_id_by_path,
    )


############################
# SyncKnowledgeCleanup
############################


class SyncCleanupForm(BaseModel):
    file_ids: list[str]  # file IDs to delete
    dir_ids: list[str] = []  # directory IDs to rmdir


@router.post('/{id}/sync/cleanup')
async def sync_knowledge_cleanup(
    id: str,
    form_data: SyncCleanupForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """
    Remove stale files and orphaned directories from a knowledge base
    after an incremental sync.
    """
    await _verify_knowledge_write_access(id, user, db)

    # ── Remove deleted files ──
    for file_id in form_data.file_ids:
        file = await Files.get_file_by_id(file_id, db=db)
        if not file:
            continue

        await Knowledges.remove_file_from_knowledge_by_id(id, file_id, db=db)

        try:
            await ASYNC_VECTOR_DB_CLIENT.delete(collection_name=id, filter={'file_id': file_id})
            await ASYNC_VECTOR_DB_CLIENT.delete(collection_name=id, filter={'hash': file.hash})
        except Exception:
            pass

        try:
            collection_name = f'file-{file_id}'
            if await ASYNC_VECTOR_DB_CLIENT.has_collection(collection_name):
                await ASYNC_VECTOR_DB_CLIENT.delete_collection(collection_name)
        except Exception:
            pass

        if file.user_id == user.id or user.role == 'admin':
            await Files.delete_file_by_id(file_id, db=db)
            try:
                await asyncio.to_thread(Storage.delete_file, file.path)
            except Exception:
                pass

    # ── Remove orphaned directories (children before parents) ──
    for dir_id in reversed(form_data.dir_ids):
        await Knowledges.delete_directory(dir_id, move_files_to_parent=False, db=db)

    return {'status': True}


############################
# AddFilesToKnowledge
############################


@router.post('/{id}/files/batch/add', response_model=KnowledgeFilesResponse | None)
async def add_files_to_knowledge_batch(
    request: Request,
    id: str,
    form_data: list[KnowledgeFileIdForm],
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """
    Add multiple files to a knowledge base
    """
    knowledge = await Knowledges.get_knowledge_by_id(id=id, db=db)
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )
    if is_external_knowledge(knowledge):
        external_knowledge_error()

    if (
        knowledge.user_id != user.id
        and not await AccessGrants.has_access(
            user_id=user.id,
            resource_type='knowledge',
            resource_id=knowledge.id,
            permission='write',
            db=db,
        )
        and user.role != 'admin'
    ):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.ACCESS_PROHIBITED,
        )

    # Batch-fetch all files to avoid N+1 queries
    log.info(f'files/batch/add - {len(form_data)} files')
    file_ids = [form.file_id for form in form_data]
    files = await Files.get_files_by_ids(file_ids, db=db)

    # Verify all requested files were found
    found_ids = {file.id for file in files}
    missing_ids = [fid for fid in file_ids if fid not in found_ids]
    if missing_ids:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f'File {missing_ids[0]} not found',
        )

    # Per-file read-access check — same gate as the single-file endpoint.
    if user.role != 'admin':
        for file in files:
            if file.user_id != user.id and not await has_access_to_file(file.id, 'read', user, db=db):
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail=ERROR_MESSAGES.ACCESS_PROHIBITED,
                )

    # Filter out files already linked to this knowledge base to prevent
    # duplicate embeddings in the vector DB (issue #10679).
    new_entries = []
    for form in form_data:
        if not await Knowledges.has_file(knowledge_id=id, file_id=form.file_id, db=db):
            new_entries.append(form)

    if not new_entries:
        return KnowledgeFilesResponse(
            **knowledge.model_dump(),
            files=await Knowledges.get_file_metadatas_by_id(knowledge.id, db=db),
        )

    # Narrow the file list to only new files for processing
    new_file_ids = {form.file_id for form in new_entries}
    files = [f for f in files if f.id in new_file_ids]

    # Process files
    try:
        result = await process_files_batch(
            request=request,
            form_data=BatchProcessFilesForm(files=files, collection_name=id),
            user=user,
            db=db,
        )
    except Exception as e:
        log.error(f'add_files_to_knowledge_batch: Exception occurred: {e}', exc_info=True)
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))

    # Only add files that were successfully processed
    successful_file_ids = [r.file_id for r in result.results if r.status == 'completed']
    dir_map = {form.file_id: form.directory_id for form in new_entries}
    for file_id in successful_file_ids:
        await Knowledges.add_file_to_knowledge_by_id(
            knowledge_id=id,
            file_id=file_id,
            user_id=user.id,
            directory_id=dir_map.get(file_id),
            db=db,
        )

    # If there were any errors, include them in the response
    if result.errors:
        error_details = [f'{err.file_id}: {err.error}' for err in result.errors]
        return KnowledgeFilesResponse(
            **knowledge.model_dump(),
            files=await Knowledges.get_file_metadatas_by_id(knowledge.id, db=db),
            warnings={
                'message': 'Some files failed to process',
                'errors': error_details,
            },
        )

    return KnowledgeFilesResponse(
        **knowledge.model_dump(),
        files=await Knowledges.get_file_metadatas_by_id(knowledge.id, db=db),
    )


############################
# ExportKnowledgeById
############################


@router.get('/{id}/export')
async def export_knowledge_by_id(id: str, user=Depends(get_admin_user), db: AsyncSession = Depends(get_async_session)):
    """
    Export a knowledge base as a zip file containing .txt files.
    Admin only.
    """

    knowledge = await Knowledges.get_knowledge_by_id(id=id, db=db)
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )
    if is_external_knowledge(knowledge):
        external_knowledge_error()

    files = await Knowledges.get_files_by_id(id, db=db)

    # Create zip file in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for file in files:
            content = file.data.get('content', '') if file.data else ''
            if content:
                # Use original filename with .txt extension
                filename = file.filename
                if not filename.endswith('.txt'):
                    filename = f'{filename}.txt'
                zf.writestr(filename, content)

    zip_buffer.seek(0)

    # Sanitize knowledge name for filename
    safe_name = ''.join(c if c.isalnum() or c in ' -_' else '_' for c in knowledge.name)
    zip_filename = f'{safe_name}.zip'

    return StreamingResponse(
        zip_buffer,
        media_type='application/zip',
        headers={'Content-Disposition': f"attachment; filename*=UTF-8''{quote(zip_filename, safe='')}"},
    )


############################
# Directory endpoints
############################


class KnowledgeDirectoryCreateForm(BaseModel):
    name: str
    parent_id: Optional[str] = None


class KnowledgeDirectoryUpdateForm(BaseModel):
    name: Optional[str] = None
    parent_id: Optional[str] = '__unset__'


class KnowledgeFileMoveForm(BaseModel):
    file_id: str
    directory_id: Optional[str] = None


async def _verify_knowledge_write_access(id: str, user, db: AsyncSession):
    """Verify the user has write access to the knowledge base. Returns the knowledge model."""
    knowledge = await Knowledges.get_knowledge_by_id(id=id, db=db)
    if not knowledge:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )
    if is_external_knowledge(knowledge):
        external_knowledge_error()
    if (
        knowledge.user_id != user.id
        and not await AccessGrants.has_access(
            user_id=user.id,
            resource_type='knowledge',
            resource_id=knowledge.id,
            permission='write',
            db=db,
        )
        and user.role != 'admin'
    ):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail=ERROR_MESSAGES.ACCESS_PROHIBITED,
        )
    return knowledge


@router.post('/{id}/dirs/create', response_model=KnowledgeDirectoryModel)
async def create_knowledge_directory(
    request: Request,
    id: str,
    form_data: KnowledgeDirectoryCreateForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    await _verify_knowledge_write_access(id, user, db)

    directory = await Knowledges.create_directory(
        knowledge_id=id,
        name=form_data.name,
        user_id=user.id,
        parent_id=form_data.parent_id,
        db=db,
    )
    if not directory:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail='Failed to create directory. A directory with this name may already exist at this level.',
        )
    await publish_event(
        request,
        EVENTS.KNOWLEDGE_DIRECTORY_CREATED,
        actor=user,
        subject_id=directory.id,
        data={'knowledge_id': id, 'name': directory.name, 'parent_id': directory.parent_id},
    )
    return directory


@router.post('/{id}/dirs/{dir_id}/update', response_model=KnowledgeDirectoryModel)
async def update_knowledge_directory(
    request: Request,
    id: str,
    dir_id: str,
    form_data: KnowledgeDirectoryUpdateForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    await _verify_knowledge_write_access(id, user, db)

    # Verify directory belongs to this knowledge base
    directory = await Knowledges.get_directory_by_id(dir_id, db=db)
    if not directory or directory.knowledge_id != id:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    result = await Knowledges.update_directory(
        directory_id=dir_id,
        name=form_data.name,
        parent_id=form_data.parent_id,
        db=db,
    )
    if not result:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail='Failed to update directory. This may be caused by a naming conflict or circular move.',
        )
    await publish_event(
        request,
        EVENTS.KNOWLEDGE_DIRECTORY_UPDATED,
        actor=user,
        subject_id=result.id,
        data={'knowledge_id': id, 'name': result.name, 'parent_id': result.parent_id},
    )
    return result


@router.delete('/{id}/dirs/{dir_id}/delete')
async def delete_knowledge_directory(
    request: Request,
    id: str,
    dir_id: str,
    move_files: bool = Query(True, description='If true, move contained files to parent. If false, delete them.'),
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    await _verify_knowledge_write_access(id, user, db)

    # Verify directory belongs to this knowledge base
    directory = await Knowledges.get_directory_by_id(dir_id, db=db)
    if not directory or directory.knowledge_id != id:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    success = await Knowledges.delete_directory(
        directory_id=dir_id,
        move_files_to_parent=move_files,
        db=db,
    )
    if not success:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail='Failed to delete directory.',
        )
    await publish_event(
        request,
        EVENTS.KNOWLEDGE_DIRECTORY_DELETED,
        actor=user,
        subject_id=dir_id,
        data={'knowledge_id': id, 'move_files': move_files},
    )
    return {'status': True}


@router.post('/{id}/file/move')
async def move_file_in_knowledge(
    request: Request,
    id: str,
    form_data: KnowledgeFileMoveForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    await _verify_knowledge_write_access(id, user, db)

    # Verify file belongs to this knowledge base
    if not await Knowledges.has_file(knowledge_id=id, file_id=form_data.file_id, db=db):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    # If target directory is set, verify it belongs to this knowledge base
    if form_data.directory_id:
        directory = await Knowledges.get_directory_by_id(form_data.directory_id, db=db)
        if not directory or directory.knowledge_id != id:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail='Target directory not found.',
            )

    success = await Knowledges.move_file_to_directory(
        knowledge_id=id,
        file_id=form_data.file_id,
        directory_id=form_data.directory_id,
        db=db,
    )
    if not success:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail='Failed to move file.',
        )
    await publish_event(
        request,
        EVENTS.KNOWLEDGE_FILE_MOVED,
        actor=user,
        subject_id=form_data.file_id,
        data={'knowledge_id': id, 'directory_id': form_data.directory_id},
    )
    return {'status': True}


# ─────────────────────────────────────────────────────────────
# Enterprise Knowledge Dashboard – Phase 1: Chunk Management
# ─────────────────────────────────────────────────────────────


@router.post('/{id}/chunks/preview', response_model=list[KnowledgeChunkModel])
async def preview_knowledge_chunks(
    request: Request,
    id: str,
    form_data: KnowledgeChunkPreviewForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Chunk a file without embedding – preview what the splitter will produce."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    file = await Files.get_file_by_id(form_data.file_id)
    if not file:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    file_path = file.path
    if not file_path:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='File has no path.')

    # Use existing loader + splitter (no embedding)
    try:
        file_path = await asyncio.to_thread(Storage.get_file, file_path)

        loader_config = await get_loader_config()
        loader = build_loader_from_config(request, loader_config)
        loader.user = user
        loader.metadata = {
            'file_id': file.id,
            'file_name': file.filename,
            'file_content_type': file.meta.get('content_type'),
        }
        docs = await loader.aload(file.filename, file.meta.get('content_type'), file_path)

        method = getattr(form_data, 'method', None) or 'general'
        if method and method not in ('', 'general'):
            from open_webui.utils.chunking_strategies import chunk_document, extract_keywords, generate_questions
            # Use loader-extracted text (handles txt/pdf/docx/etc.); custom strategies
            # fall through to the default splitter below when no text can be extracted.
            content = '\n\n'.join(doc.page_content for doc in docs if doc.page_content)
            if content:
                raw = chunk_document(content, method)
                from sqlalchemy import delete as sa_delete
                await db.execute(sa_delete(KnowledgeChunk).where(KnowledgeChunk.knowledge_id == id, KnowledgeChunk.file_id == form_data.file_id))
                import hashlib
                now = int(time.time())
                chunks = []
                for idx, rc in enumerate(raw):
                    kw = extract_keywords(rc['content'])
                    qs = generate_questions(rc['content'])
                    h = hashlib.sha256(rc['content'].encode()).hexdigest()
                    ck = KnowledgeChunk(id=str(uuid.uuid4()), knowledge_id=id, file_id=form_data.file_id, chunk_index=idx, content=rc['content'], token_count=len(rc['content']) // 4, content_hash=h, meta={**rc.get('metadata', {}), 'keywords': kw, 'questions': qs}, created_at=now, updated_at=now)
                    db.add(ck)
                    chunks.append(KnowledgeChunkModel.model_validate(ck))
                await db.commit()
                return chunks

        # Chunk using same logic as save_docs_to_vector_db
        config = await get_retrieval_config()
        from open_webui.routers.retrieval import merge_docs_to_target_size
        from langchain_text_splitters import (
            MarkdownHeaderTextSplitter,
            RecursiveCharacterTextSplitter,
            TokenTextSplitter,
        )
        import tiktoken

        if config.ENABLE_MARKDOWN_HEADER_TEXT_SPLITTER:
            markdown_splitter = MarkdownHeaderTextSplitter(
                headers_to_split_on=[
                    ('#', 'Header 1'),
                    ('##', 'Header 2'),
                    ('###', 'Header 3'),
                    ('####', 'Header 4'),
                    ('#####', 'Header 5'),
                    ('######', 'Header 6'),
                ],
                strip_headers=False,
            )
            split_docs = []
            for doc in docs:
                split_docs.extend([
                    type(doc)(page_content=c.page_content, metadata={**doc.metadata})
                    for c in markdown_splitter.split_text(doc.page_content)
                ])
            docs = split_docs
            if config.CHUNK_MIN_SIZE_TARGET > 0:
                docs = merge_docs_to_target_size(request, docs, config)

        if config.TEXT_SPLITTER in ['', 'character']:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=config.CHUNK_SIZE, chunk_overlap=config.CHUNK_OVERLAP, add_start_index=True,
            )
            docs = text_splitter.split_documents(docs)
        elif config.TEXT_SPLITTER == 'token':
            tiktoken.get_encoding(str(config.TIKTOKEN_ENCODING_NAME))
            text_splitter = TokenTextSplitter(
                encoding_name=str(config.TIKTOKEN_ENCODING_NAME),
                chunk_size=config.CHUNK_SIZE, chunk_overlap=config.CHUNK_OVERLAP, add_start_index=True,
            )
            docs = text_splitter.split_documents(docs)
        else:
            # Fallback to RecursiveCharacterTextSplitter
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=config.CHUNK_SIZE, chunk_overlap=config.CHUNK_OVERLAP, add_start_index=True,
            )
            docs = text_splitter.split_documents(docs)

        # Remove old preview chunks for this file in this KB
        from sqlalchemy import delete as sa_delete
        await db.execute(
            sa_delete(KnowledgeChunk).where(
                KnowledgeChunk.knowledge_id == id,
                KnowledgeChunk.file_id == form_data.file_id,
            )
        )

        # Persist chunks in knowledge_chunk table
        import hashlib
        from open_webui.utils.chunking_strategies import extract_keywords, generate_questions
        now = int(time.time())
        chunks = []
        for idx, doc in enumerate(docs):
            content_hash = hashlib.sha256(doc.page_content.encode()).hexdigest()
            token_count = len(doc.page_content) // 4  # rough estimate

            chunk = KnowledgeChunk(
                id=str(uuid.uuid4()),
                knowledge_id=id,
                file_id=form_data.file_id,
                chunk_index=idx,
                content=doc.page_content,
                token_count=token_count,
                meta={**doc.metadata, 'keywords': extract_keywords(doc.page_content), 'questions': generate_questions(doc.page_content)},
                content_hash=content_hash,
                created_at=now,
                updated_at=now,
            )
            db.add(chunk)
            chunks.append(KnowledgeChunkModel.model_validate(chunk))

        await db.commit()
        return chunks
    except Exception as e:
        log.exception(e)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))


@router.get('/{id}/files/{file_id}/chunks', response_model=list[KnowledgeChunkModel])
async def get_file_chunks(
    id: str,
    file_id: str,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Get all chunks for a file in a knowledge base."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    from sqlalchemy import select as sa_select
    result = await db.execute(
        sa_select(KnowledgeChunk)
        .where(KnowledgeChunk.knowledge_id == id, KnowledgeChunk.file_id == file_id)
        .order_by(KnowledgeChunk.chunk_index.asc())
    )
    rows = result.scalars().all()
    return [KnowledgeChunkModel.model_validate(r) for r in rows]


@router.get('/{id}/chunks/{chunk_id}', response_model=KnowledgeChunkModel)
async def get_chunk_detail(
    id: str,
    chunk_id: str,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Get full content of a single chunk."""
    from sqlalchemy import select as sa_select
    result = await db.execute(
        sa_select(KnowledgeChunk).where(
            KnowledgeChunk.knowledge_id == id,
            KnowledgeChunk.id == chunk_id,
        )
    )
    chunk = result.scalars().first()
    if not chunk:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail='Chunk not found.')
    return KnowledgeChunkModel.model_validate(chunk)


@router.post('/{id}/chunks/merge')
async def merge_knowledge_chunks(
    request: Request,
    id: str,
    form_data: KnowledgeChunkMergeForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Merge adjacent chunks [start_index, end_index] into one."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    from sqlalchemy import select as sa_select, delete as sa_delete

    # Fetch chunks in range
    result = await db.execute(
        sa_select(KnowledgeChunk)
        .where(
            KnowledgeChunk.knowledge_id == id,
            KnowledgeChunk.file_id == form_data.file_id,
            KnowledgeChunk.chunk_index >= form_data.start_index,
            KnowledgeChunk.chunk_index <= form_data.end_index,
        )
        .order_by(KnowledgeChunk.chunk_index.asc())
    )
    target_chunks = result.scalars().all()
    if len(target_chunks) < 2:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Need at least 2 chunks to merge.')

    # Merge content
    merged_content = '\n\n'.join(c.content for c in target_chunks)
    merged_meta = target_chunks[0].meta or {}

    import hashlib
    now = int(time.time())
    content_hash = hashlib.sha256(merged_content.encode()).hexdigest()
    token_count = len(merged_content) // 4

    # Create new merged chunk at the first index
    merged_chunk = KnowledgeChunk(
        id=str(uuid.uuid4()),
        knowledge_id=id,
        file_id=form_data.file_id,
        chunk_index=form_data.start_index,
        content=merged_content,
        token_count=token_count,
        meta=merged_meta,
        content_hash=content_hash,
        created_at=now,
        updated_at=now,
    )

    # Delete old chunks
    old_ids = [c.id for c in target_chunks]
    await db.execute(sa_delete(KnowledgeChunk).where(KnowledgeChunk.id.in_(old_ids)))
    db.add(merged_chunk)

    # Re-index chunks after the deleted range
    result = await db.execute(
        sa_select(KnowledgeChunk)
        .where(
            KnowledgeChunk.knowledge_id == id,
            KnowledgeChunk.file_id == form_data.file_id,
            KnowledgeChunk.chunk_index > form_data.end_index,
        )
        .order_by(KnowledgeChunk.chunk_index.asc())
    )
    trailing = result.scalars().all()
    offset = form_data.end_index - form_data.start_index
    for chunk in trailing:
        chunk.chunk_index -= offset
        chunk.updated_at = now

    await db.commit()
    return {'status': True, 'merged_chunk_id': merged_chunk.id}


@router.post('/{id}/chunks/split')
async def split_knowledge_chunk(
    request: Request,
    id: str,
    form_data: KnowledgeChunkSplitForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Split a chunk at a character offset into two chunks."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    from sqlalchemy import select as sa_select

    result = await db.execute(
        sa_select(KnowledgeChunk).where(
            KnowledgeChunk.knowledge_id == id,
            KnowledgeChunk.id == form_data.chunk_id,
        )
    )
    chunk = result.scalars().first()
    if not chunk:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail='Chunk not found.')

    content = chunk.content
    split_point = form_data.split_at
    if split_point <= 0 or split_point >= len(content):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Split point must be within chunk content.')

    part_a = content[:split_point].strip()
    part_b = content[split_point:].strip()
    if not part_a or not part_b:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail='Both parts must be non-empty.')

    import hashlib
    now = int(time.time())

    # Update original chunk to part_a
    chunk.content = part_a
    chunk.token_count = len(part_a) // 4
    chunk.content_hash = hashlib.sha256(part_a.encode()).hexdigest()
    chunk.updated_at = now

    # Create new chunk for part_b
    new_chunk = KnowledgeChunk(
        id=str(uuid.uuid4()),
        knowledge_id=id,
        file_id=chunk.file_id,
        chunk_index=chunk.chunk_index + 1,
        content=part_b,
        token_count=len(part_b) // 4,
        meta=(chunk.meta or {}).copy(),
        content_hash=hashlib.sha256(part_b.encode()).hexdigest(),
        created_at=now,
        updated_at=now,
    )
    db.add(new_chunk)

    # Shift trailing chunks
    result = await db.execute(
        sa_select(KnowledgeChunk)
        .where(
            KnowledgeChunk.knowledge_id == id,
            KnowledgeChunk.file_id == chunk.file_id,
            KnowledgeChunk.chunk_index > chunk.chunk_index,
        )
        .order_by(KnowledgeChunk.chunk_index.asc())
    )
    trailing = result.scalars().all()
    for tc in trailing:
        tc.chunk_index += 1
        tc.updated_at = now

    await db.commit()
    return {'status': True, 'new_chunk_id': new_chunk.id}


@router.post('/{id}/chunks/reindex')
async def reindex_knowledge_chunks(
    request: Request,
    id: str,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Re-embed all chunks for this KB after manual adjustments."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    from sqlalchemy import select as sa_select

    result = await db.execute(
        sa_select(KnowledgeChunk).where(KnowledgeChunk.knowledge_id == id)
    )
    chunks = result.scalars().all()
    if not chunks:
        return {'status': True, 'message': 'No chunks to reindex.'}

    # Delete old vectors for this KB's collection
    try:
        await ASYNC_VECTOR_DB_CLIENT.delete_collection(collection_name=id)
    except Exception:
        pass

    # Rebuild vectors from knowledge_chunk rows
    from langchain_core.documents import Document

    docs = [
        Document(
            page_content=c.content,
            metadata={
                **(c.meta or {}),
                'file_id': c.file_id,
                'chunk_index': c.chunk_index,
                'content_hash': c.content_hash,
            },
        )
        for c in chunks
    ]

    try:
        from fastapi.concurrency import run_in_threadpool
        config = await get_retrieval_config()
        await run_in_threadpool(
            save_docs_to_vector_db,
            request, docs, id, config,
            None, True, True, False, user,
        )
    except Exception as e:
        log.exception(e)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))

    return {'status': True, 'chunks_processed': len(chunks)}


# ─────────────────────────────────────────────────────────────
# Enterprise Knowledge Dashboard – Phase 2: Progress Tracking
# ─────────────────────────────────────────────────────────────


# In-memory progress store for real-time SSE updates
# Key: (knowledge_id, file_id) → KnowledgeProcessingTaskModel
_progress_store: dict[str, KnowledgeProcessingTaskModel] = {}


async def _update_processing_progress(
    knowledge_id: str,
    file_id: str,
    status: str,
    progress_pct: int = 0,
    chunks_total: int | None = None,
    chunks_processed: int | None = None,
    error_message: str | None = None,
    db: AsyncSession | None = None,
):
    """Create or update a processing task record in DB and in-memory store."""
    import time as _time, uuid as _uuid

    key = f'{knowledge_id}:{file_id}'
    now = int(_time.time())

    # Try to update existing task
    from sqlalchemy import select as sa_select
    result = await db.execute(
        sa_select(KnowledgeProcessingTask).where(
            KnowledgeProcessingTask.knowledge_id == knowledge_id,
            KnowledgeProcessingTask.file_id == file_id,
        )
    )
    task = result.scalars().first()

    if task:
        task.status = status
        task.progress_pct = progress_pct
        if chunks_total is not None:
            task.chunks_total = chunks_total
        if chunks_processed is not None:
            task.chunks_processed = chunks_processed
        if error_message is not None:
            task.error_message = error_message
        task.updated_at = now
        db.add(task)
        await db.commit()
    else:
        task = KnowledgeProcessingTask(
            id=str(_uuid.uuid4()),
            knowledge_id=knowledge_id,
            file_id=file_id,
            task_type='full_process',
            status=status,
            progress_pct=progress_pct,
            chunks_total=chunks_total,
            chunks_processed=chunks_processed,
            error_message=error_message,
            created_at=now,
            updated_at=now,
        )
        db.add(task)
        await db.commit()

    # Update in-memory store
    _progress_store[key] = KnowledgeProcessingTaskModel.model_validate(task)


@router.get('/{id}/progress', response_model=list[KnowledgeProcessingTaskModel])
async def get_processing_progress(
    id: str,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Get current processing status for all files in a KB."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    from sqlalchemy import select as sa_select

    result = await db.execute(
        sa_select(KnowledgeProcessingTask)
        .where(KnowledgeProcessingTask.knowledge_id == id)
        .order_by(KnowledgeProcessingTask.created_at.desc())
    )
    tasks = result.scalars().all()
    return [KnowledgeProcessingTaskModel.model_validate(t) for t in tasks]


@router.get('/{id}/progress/stream')
async def stream_processing_progress(
    request: Request,
    id: str,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """SSE endpoint that streams progress updates every second."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    async def event_generator():
        import time as _time
        from sqlalchemy import select as sa_select

        # Track the last run to stop after idle timeout
        last_active = _time.time()
        sent_count = 0

        while True:
            # Check DB for tasks
            result = await db.execute(
                sa_select(KnowledgeProcessingTask)
                .where(KnowledgeProcessingTask.knowledge_id == id)
                .order_by(KnowledgeProcessingTask.created_at.desc())
            )
            tasks = result.scalars().all()
            task_list = [KnowledgeProcessingTaskModel.model_validate(t).model_dump() for t in tasks]

            # Also include in-memory tasks that may not be in DB yet
            for key, mem_task in _progress_store.items():
                if key.startswith(f'{id}:'):
                    exists = any(t['id'] == mem_task.id for t in task_list)
                    if not exists:
                        task_list.append(mem_task.model_dump())

            data = json.dumps(task_list)
            yield f'data: {data}\n\n'
            sent_count += 1

            # Check if all tasks are in terminal state
            all_done = all(
                t['status'] in ('completed', 'failed') for t in task_list
            ) if task_list else False

            if all_done and sent_count > 2:
                # Send one more update then stop
                yield f'data: {data}\n\n'
                break

            if not all_done:
                last_active = _time.time()

            # Timeout after 5 minutes of inactivity
            if _time.time() - last_active > 300:
                break

            await asyncio.sleep(1)

            # Refresh DB session between iterations
            await db.commit()

    return StreamingResponse(event_generator(), media_type='text/event-stream')


@router.get('/{id}/progress/batch', response_model=KnowledgeBatchTaskModel | None)
async def get_batch_progress(
    id: str,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Get the latest batch processing task for this KB."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    from sqlalchemy import select as sa_select

    result = await db.execute(
        sa_select(KnowledgeBatchTask)
        .where(KnowledgeBatchTask.knowledge_id == id)
        .order_by(KnowledgeBatchTask.created_at.desc())
        .limit(1)
    )
    batch = result.scalars().first()
    if not batch:
        return None
    return KnowledgeBatchTaskModel.model_validate(batch)


# ─────────────────────────────────────────────────────────────
# Enterprise Knowledge Dashboard – Phase 3: Retrieval Evaluation
# ─────────────────────────────────────────────────────────────


@router.post('/{id}/evaluate/query')
async def evaluate_retrieval_query(
    request: Request,
    id: str,
    form_data: KnowledgeEvaluateQueryForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Run a test query against the KB and return Top-K results with metrics."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    k = form_data.k if form_data.k else 10

    try:
        result = await query_collection(
            request=request,
            collection_names=[id],
            queries=[form_data.query],
            embedding_function=request.app.state.EMBEDDING_FUNCTION,
            k=k,
        )
    except Exception as e:
        log.exception(e)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))

    # query_collection returns: {'distances': [[]], 'documents': [[]], 'metadatas': [[]]}
    results = []
    if isinstance(result, dict):
        docs_list = result.get('documents', [[]])[0] if result.get('documents') else []
        metas_list = result.get('metadatas', [[]])[0] if result.get('metadatas') else []
        dists_list = result.get('distances', [[]])[0] if result.get('distances') else []

        for i in range(min(k, len(docs_list))):
            doc_text = docs_list[i] if i < len(docs_list) else ''
            results.append({
                'chunk_id': f'result-{i}-' + (metas_list[i].get('content_hash', '')[:8] if i < len(metas_list) else ''),
                'text': doc_text[:500],
                'score': dists_list[i] if i < len(dists_list) else None,
                'metadata': metas_list[i] if i < len(metas_list) else {},
                'rank': i + 1,
            })

    # Compute metrics if judgments exist
    from sqlalchemy import select as sa_select
    metrics = None
    j_result = await db.execute(
        sa_select(KnowledgeRelevanceJudgment)
        .where(
            KnowledgeRelevanceJudgment.knowledge_id == id,
            KnowledgeRelevanceJudgment.query_text == form_data.query,
        )
    )
    judgments = j_result.scalars().all()
    if judgments:
        relevant_ids = {j.chunk_id for j in judgments if j.relevance == 1}
        total_relevant = len(relevant_ids)
        if total_relevant > 0:
            # recall@K
            retrieved_ids = {r['chunk_id'] for r in results if r['chunk_id']}
            relevant_retrieved = relevant_ids & retrieved_ids
            recall_at_k = len(relevant_retrieved) / total_relevant
            # precision@K
            precision_at_k = len(relevant_retrieved) / k if k > 0 else 0
            # MRR
            mrr = 0.0
            for j in judgments:
                if j.relevance == 1 and j.rank_position and j.rank_position > 0:
                    mrr = max(mrr, 1.0 / j.rank_position)
            metrics = {
                'recall_at_k': round(recall_at_k, 4),
                'precision_at_k': round(precision_at_k, 4),
                'mrr': round(mrr, 4),
                'total_relevant': total_relevant,
            }

    return {'results': results, 'metrics': metrics}


@router.post('/{id}/evaluate/annotate')
async def annotate_retrieval_results(
    request: Request,
    id: str,
    form_data: KnowledgeRelevanceAnnotationForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Mark retrieved chunks as relevant (1) or not relevant (0)."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    from sqlalchemy import delete as sa_delete
    now = int(time.time())
    count = 0

    for j in form_data.judgments:
        # Upsert: delete old judgment for same query+chunk, then insert
        await db.execute(
            sa_delete(KnowledgeRelevanceJudgment).where(
                KnowledgeRelevanceJudgment.knowledge_id == id,
                KnowledgeRelevanceJudgment.query_text == form_data.query_text,
                KnowledgeRelevanceJudgment.chunk_id == j.get('chunk_id'),
            )
        )
        judgment = KnowledgeRelevanceJudgment(
            id=str(uuid.uuid4()),
            knowledge_id=id,
            query_text=form_data.query_text,
            chunk_id=j.get('chunk_id'),
            document_text=j.get('document_text', ''),
            rank_position=j.get('rank_position'),
            relevance=j.get('relevance', 0),
            user_id=user.id,
            created_at=now,
        )
        db.add(judgment)
        count += 1

    await db.commit()
    return {'status': True, 'annotations_saved': count}


@router.get('/{id}/evaluate/judgments')
async def get_evaluation_judgments(
    id: str,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Get all relevance judgments for this KB, grouped by query."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    from sqlalchemy import select as sa_select
    result = await db.execute(
        sa_select(KnowledgeRelevanceJudgment)
        .where(KnowledgeRelevanceJudgment.knowledge_id == id)
        .order_by(KnowledgeRelevanceJudgment.created_at.desc())
    )
    rows = result.scalars().all()

    # Group by query_text
    grouped: dict[str, list] = {}
    for r in rows:
        key = r.query_text
        if key not in grouped:
            grouped[key] = []
        grouped[key].append(KnowledgeRelevanceJudgmentModel.model_validate(r).model_dump())

    return {'judgments': grouped, 'total_queries': len(grouped)}


@router.delete('/{id}/evaluate/judgments/{query_text:path}')
async def delete_evaluation_judgments(
    id: str,
    query_text: str,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Delete all judgments for a specific query."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    from urllib.parse import unquote
    from sqlalchemy import delete as sa_delete

    decoded_query = unquote(query_text)
    await db.execute(
        sa_delete(KnowledgeRelevanceJudgment).where(
            KnowledgeRelevanceJudgment.knowledge_id == id,
            KnowledgeRelevanceJudgment.query_text == decoded_query,
        )
    )
    await db.commit()
    return {'status': True}


# ─────────────────────────────────────────────────────────────
# Enterprise Knowledge Dashboard – Phase 11: Faithfulness (LLM-as-judge)
# ─────────────────────────────────────────────────────────────

_JUDGE_SYSTEM = (
    "You are a strict fact-checker evaluating the faithfulness of a RAG answer. "
    "Judge whether each claim in the answer is supported by the provided context. "
    "Use ONLY the context, never outside knowledge. Output ONLY a valid JSON object, no markdown fences."
)

_JUDGE_PROMPT = """Given a question, the retrieved context (inside <source> tags), and a generated answer:

1. Decompose the answer into a list of atomic claims (each is a single verifiable statement).
2. For each claim, set "supported":
   - "yes"         → the context directly supports the claim;
   - "no"          → the context contradicts the claim;
   - "insufficient" → the context does not mention enough to verify the claim.
3. If the answer is a refusal (e.g. says the information is not found), output an empty claims list.

Output a JSON object exactly in this shape:
{{"claims": [{{"text": "...", "supported": "yes|no|insufficient", "evidence": "..."}}]}}

Question: {question}

Context:
{context}

Answer:
{answer}
"""


def _faithfulness_api() -> tuple[str, str, str]:
    """返回 (base_url, api_key, model)。优先 .env 直读（对齐 exec_workflow），回退全局配置，默认 DeepSeek。"""
    # config.* 常量在模块导入时取值，可能早于 env.py 加载 .env 而为空；
    # 这里与 exec_workflow 一致：显式 load_dotenv(override=True) 后直读 os.getenv。
    try:
        from dotenv import load_dotenv

        load_dotenv(override=True)
    except ImportError:
        pass
    base = (
        os.getenv("RAG_OPENAI_API_BASE_URL")
        or os.getenv("OPENAI_API_BASE_URL")
        or RAG_OPENAI_API_BASE_URL
        or OPENAI_API_BASE_URL
        or "https://api.deepseek.com/v1"
    ).rstrip("/")
    key = os.getenv("RAG_OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY") or RAG_OPENAI_API_KEY or OPENAI_API_KEY or ""
    model = os.getenv("EVAL_MODEL", "deepseek-chat")
    return base, key, model


def _build_source_context(chunks: list[dict]) -> str:
    """把检索片段拼成 <source> 块，id 从 1 递增，对齐离线 eval 的 build_context_string。"""
    parts = []
    for i, chunk in enumerate(chunks, start=1):
        text = chunk.get("text", "") or chunk.get("content", "")
        src_name = chunk.get("source") or chunk.get("name") or ""
        attrs = f' name="{src_name}"' if src_name else ""
        parts.append(f'<source id="{i}"{attrs}>{text}</source>')
    return "\n".join(parts)


def _render_rag_prompt(question: str, chunks: list[dict]) -> str:
    """渲染 RAG 模板：注入 context 与 query（对齐离线 eval 的 render_rag_prompt）。"""
    context = _build_source_context(chunks)
    template = RAG_TEMPLATE or ""
    prompt = template.replace("{{CONTEXT}}", context).replace("[context]", context)
    prompt = prompt.replace("{{QUERY}}", question).replace("[query]", question)
    return prompt


def _extract_json(text: str) -> dict:
    """从 LLM 输出稳健抽取 JSON（容忍 markdown 围栏 / 前后缀）。"""
    text = (text or "").strip()
    fence = re.search(r"```(?:json)?\s*(.*?)```", text, re.DOTALL)
    if fence:
        text = fence.group(1).strip()
    start, end = text.find("{"), text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError(f"no JSON object in judge output: {text[:200]}")
    return json.loads(text[start : end + 1])


async def _faithfulness_chat(messages: list[dict], temperature: float = 0.0) -> str:
    """调 DeepSeek（OpenAI 兼容协议）拿一段文本输出。"""
    base, key, model = _faithfulness_api()
    if not key:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No API key configured for faithfulness evaluation (set OPENAI_API_KEY or RAG_OPENAI_API_KEY)",
        )
    async with httpx.AsyncClient(timeout=120) as client:
        resp = await client.post(
            base + "/chat/completions",
            headers={"Authorization": f"Bearer {key}"},
            json={"model": model, "messages": messages, "temperature": temperature, "stream": False},
        )
        if resp.status_code != 200:
            raise HTTPException(status_code=status.HTTP_502_BAD_GATEWAY, detail=f"LLM call failed: HTTP {resp.status_code}")
        data = resp.json()
        return (data.get("choices", [{}])[0].get("message", {}).get("content") or "").strip()


async def _generate_answer(question: str, chunks: list[dict]) -> str:
    """给定问题 + 检索片段，用 RAG 模板生成答案（system=渲染后的模板, user=问题）。"""
    return await _faithfulness_chat(
        [
            {"role": "system", "content": _render_rag_prompt(question, chunks)},
            {"role": "user", "content": question},
        ]
    )


async def _judge_faithfulness(question: str, answer: str, chunks: list[dict]) -> dict:
    """LLM-as-judge 忠实度判定：返回 faithfulness / claims / supported / total / refusal。"""
    if not answer.strip():
        return {"faithfulness": 0.0, "claims": [], "supported": 0, "total": 0, "refusal": False}

    context = _build_source_context(chunks)
    messages = [
        {"role": "system", "content": _JUDGE_SYSTEM},
        {"role": "user", "content": _JUDGE_PROMPT.format(question=question, context=context, answer=answer)},
    ]

    data = None
    last_err = None
    for _ in range(2):
        raw = await _faithfulness_chat(messages)
        try:
            data = _extract_json(raw)
            break
        except (ValueError, json.JSONDecodeError) as e:
            last_err = e
            messages.append({"role": "assistant", "content": raw})
            messages.append({"role": "user", "content": "That was not valid JSON. Output ONLY the JSON object as requested."})
    if data is None:
        raise HTTPException(status_code=status.HTTP_502_BAD_GATEWAY, detail=f"Judge returned invalid JSON: {last_err}")

    claims = data.get("claims", [])
    if not claims:
        # 无主张（拒答 / 空输出）→ 视为忠实，没有编造
        return {"faithfulness": 1.0, "claims": [], "supported": 0, "total": 0, "refusal": True}

    supported = sum(1 for c in claims if c.get("supported") == "yes")
    total = len(claims)
    return {
        "faithfulness": round(supported / total, 4),
        "claims": claims,
        "supported": supported,
        "total": total,
        "refusal": False,
    }


class KnowledgeFaithfulnessEvalForm(BaseModel):
    query: Optional[str] = None
    queries: Optional[List[str]] = None
    k: int = 10


@router.post('/{id}/evaluate/faithfulness')
async def evaluate_faithfulness(
    request: Request,
    id: str,
    form_data: KnowledgeFaithfulnessEvalForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """对一条或多条问题跑「真实检索 → 生成 → LLM-as-judge 忠实度判定」，返回逐条拆解 + 聚合分数。

    这是离线评测体系（enterprise-kb/eval）的 full 档在线版，供前端可视化查看生成质量。
    """
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    questions = list(form_data.queries or [])
    if form_data.query:
        questions.append(form_data.query)
    questions = [q.strip() for q in questions if q and q.strip()]
    if not questions:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="query or queries is required")

    k = form_data.k if form_data.k else 10
    _, _, model = _faithfulness_api()

    results = []
    for q in questions:
        # 1) 真实混合检索 + 重排（与 evaluate/query 同一条生产链路）
        try:
            qr = await query_collection(
                request=request,
                collection_names=[id],
                queries=[q],
                embedding_function=request.app.state.EMBEDDING_FUNCTION,
                k=k,
                k_reranker=k,
            )
        except Exception as e:
            log.exception(e)
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(e))

        chunks = []
        if isinstance(qr, dict) and qr.get("documents"):
            docs = qr.get("documents", [[]])[0]
            metas = qr.get("metadatas", [[]])[0] if qr.get("metadatas") else []
            dists = qr.get("distances", [[]])[0] if qr.get("distances") else []
            for i in range(min(k, len(docs))):
                meta = metas[i] if i < len(metas) else {}
                chunks.append({
                    "text": docs[i],
                    "source": meta.get("name") or meta.get("source") or meta.get("file_id", ""),
                    "score": dists[i] if i < len(dists) else None,
                    "rank": i + 1,
                })

        # 2) 生成答案（RAG 模板 + DeepSeek，temperature=0）
        answer = await _generate_answer(q, chunks)

        # 3) LLM-as-judge 忠实度判定
        verdict = await _judge_faithfulness(q, answer, chunks)

        results.append({
            "query": q,
            "answer": answer,
            "faithfulness": verdict["faithfulness"],
            "supported": verdict["supported"],
            "total": verdict["total"],
            "refusal": verdict["refusal"],
            "claims": verdict["claims"],
            "context_chunks": chunks,
        })

    scores = [r["faithfulness"] for r in results]
    aggregate = round(sum(scores) / len(scores), 4) if scores else 0.0
    return {
        "mode": "full",
        "model": model,
        "aggregate_faithfulness": aggregate,
        "threshold": float(os.getenv("EVAL_THRESHOLD", "0.8")),
        "results": results,
    }


# ─────────────────────────────────────────────────────────────
# Enterprise Knowledge Dashboard – Phase 4: Snapshot Management
# ─────────────────────────────────────────────────────────────


@router.post('/{id}/snapshots', response_model=KnowledgeSnapshotModel)
async def create_knowledge_snapshot(
    request: Request,
    id: str,
    form_data: KnowledgeSnapshotCreateForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Create a snapshot of the current KB state (files + chunks)."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    from sqlalchemy import select as sa_select
    now = int(time.time())

    # Collect chunk records
    chunk_result = await db.execute(
        sa_select(KnowledgeChunk).where(KnowledgeChunk.knowledge_id == id)
    )
    chunk_rows = chunk_result.scalars().all()

    # Collect file associations
    files_with_dirs = await Knowledges.get_files_with_directory_ids(id, db=db)

    # Build snapshot data
    snapshot_data = {
        'files': [
            {
                'file_id': f.id,
                'filename': f.filename,
                'directory_id': dir_id,
            }
            for f, dir_id in files_with_dirs
        ],
        'chunks': [
            {
                'id': c.id,
                'file_id': c.file_id,
                'chunk_index': c.chunk_index,
                'content_hash': c.content_hash,
                'token_count': c.token_count,
            }
            for c in chunk_rows
        ],
    }

    snapshot = KnowledgeSnapshot(
        id=str(uuid.uuid4()),
        knowledge_id=id,
        label=form_data.label,
        description=form_data.description,
        file_count=len(files_with_dirs),
        chunk_count=len(chunk_rows),
        snapshot_data=snapshot_data,
        created_by=user.id,
        created_at=now,
    )
    db.add(snapshot)
    await db.commit()
    await db.refresh(snapshot)

    return KnowledgeSnapshotModel.model_validate(snapshot)


@router.get('/{id}/snapshots', response_model=list[KnowledgeSnapshotModel])
async def list_knowledge_snapshots(
    id: str,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """List all snapshots for this KB."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    from sqlalchemy import select as sa_select
    result = await db.execute(
        sa_select(KnowledgeSnapshot)
        .where(KnowledgeSnapshot.knowledge_id == id)
        .order_by(KnowledgeSnapshot.created_at.desc())
    )
    rows = result.scalars().all()
    return [KnowledgeSnapshotModel.model_validate(r) for r in rows]


@router.get('/{id}/snapshots/{snapshot_id}', response_model=KnowledgeSnapshotModel)
async def get_knowledge_snapshot(
    id: str,
    snapshot_id: str,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Get a single snapshot detail."""
    from sqlalchemy import select as sa_select
    result = await db.execute(
        sa_select(KnowledgeSnapshot).where(
            KnowledgeSnapshot.id == snapshot_id,
            KnowledgeSnapshot.knowledge_id == id,
        )
    )
    snap = result.scalars().first()
    if not snap:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail='Snapshot not found.')
    return KnowledgeSnapshotModel.model_validate(snap)


@router.post('/{id}/snapshots/{snapshot_id}/rollback')
async def rollback_knowledge_snapshot(
    request: Request,
    id: str,
    snapshot_id: str,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Rollback to a previous snapshot - restore chunk records and re-index."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    from sqlalchemy import select as sa_select, delete as sa_delete

    result = await db.execute(
        sa_select(KnowledgeSnapshot).where(
            KnowledgeSnapshot.id == snapshot_id,
            KnowledgeSnapshot.knowledge_id == id,
        )
    )
    snap = result.scalars().first()
    if not snap:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail='Snapshot not found.')

    snap_data = snap.snapshot_data or {}
    snap_files = snap_data.get('files', [])
    snap_chunks = snap_data.get('chunks', [])

    now = int(time.time())

    # 1. Clear current chunks and file associations
    await db.execute(sa_delete(KnowledgeChunk).where(KnowledgeChunk.knowledge_id == id))
    await db.execute(sa_delete(KnowledgeFile).where(KnowledgeFile.knowledge_id == id))

    restored_chunks = len(snap_chunks)

    # 2. Restore chunk records
    for c in snap_chunks:
        db.add(KnowledgeChunk(
            id=c.get('id', str(uuid.uuid4())),
            knowledge_id=id, file_id=c['file_id'],
            chunk_index=c['chunk_index'], content='',
            token_count=c.get('token_count'), content_hash=c.get('content_hash'),
            created_at=now, updated_at=now,
        ))

    # 3. Restore file associations
    seen = set()
    for f in snap_files:
        if f['file_id'] not in seen:
            seen.add(f['file_id'])
            db.add(KnowledgeFile(
                id=str(uuid.uuid4()), knowledge_id=id,
                file_id=f['file_id'], user_id=user.id,
                directory_id=f.get('directory_id'),
                created_at=now, updated_at=now,
            ))

    await db.commit()

    return {
        'status': True,
        'restored_files': len(snap_files),
        'restored_chunks': restored_chunks,
        'snapshot_label': snap.label,
        'hint': 'Vectors need re-indexing. Go to Chunks tab and click Reindex Vectors to rebuild.',
    }


@router.post('/{id}/snapshots/compare', response_model=KnowledgeSnapshotCompareResult)
async def compare_knowledge_snapshots(
    id: str,
    form_data: KnowledgeSnapshotCompareForm,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Compare two snapshots and show added/removed/modified files."""
    knowledge = await Knowledges.get_knowledge_by_id(id, db=db)
    if not knowledge:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=ERROR_MESSAGES.NOT_FOUND)

    from sqlalchemy import select as sa_select

    # Load both snapshots
    r1 = await db.execute(
        sa_select(KnowledgeSnapshot).where(KnowledgeSnapshot.id == form_data.snapshot_a_id)
    )
    snap_a = r1.scalars().first()
    r2 = await db.execute(
        sa_select(KnowledgeSnapshot).where(KnowledgeSnapshot.id == form_data.snapshot_b_id)
    )
    snap_b = r2.scalars().first()
    if not snap_a or not snap_b:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail='Snapshot not found.')

    files_a = {(f['file_id'], f['filename']) for f in (snap_a.snapshot_data or {}).get('files', [])}
    files_b = {(f['file_id'], f['filename']) for f in (snap_b.snapshot_data or {}).get('files', [])}

    added = [{'file_id': fid, 'filename': fn} for fid, fn in (files_b - files_a)]
    removed = [{'file_id': fid, 'filename': fn} for fid, fn in (files_a - files_b)]

    # Check for modified (same file_id but different hash)
    chunks_a = {(c['file_id'], c.get('content_hash')) for c in (snap_a.snapshot_data or {}).get('chunks', [])}
    chunks_b = {(c['file_id'], c.get('content_hash')) for c in (snap_b.snapshot_data or {}).get('chunks', [])}
    common_ids = {fid for fid, _ in files_a} & {fid for fid, _ in files_b}
    modified = []
    for fid in common_ids:
        hashes_a = {h for f, h in chunks_a if f == fid}
        hashes_b = {h for f, h in chunks_b if f == fid}
        if hashes_a != hashes_b:
            name_a = next((fn for i, fn in files_a if i == fid), fid)
            modified.append({'file_id': fid, 'filename': name_a})

    return KnowledgeSnapshotCompareResult(
        added_files=added,
        removed_files=removed,
        modified_files=modified,
        total_chunks_before=snap_a.chunk_count or 0,
        total_chunks_after=snap_b.chunk_count or 0,
    )


@router.delete('/{id}/snapshots/{snapshot_id}')
async def delete_knowledge_snapshot(
    id: str,
    snapshot_id: str,
    user=Depends(get_verified_user),
    db: AsyncSession = Depends(get_async_session),
):
    """Delete a snapshot."""
    from sqlalchemy import delete as sa_delete
    await db.execute(
        sa_delete(KnowledgeSnapshot).where(
            KnowledgeSnapshot.id == snapshot_id,
            KnowledgeSnapshot.knowledge_id == id,
        )
    )
    await db.commit()
    return {'status': True}


